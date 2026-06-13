#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Morgan-Stanley-Style Research Report Generator (docx).

真实研究报告风格：品牌深蓝 ``#0B2C5C`` + 金色 ``#C8A951``，
字体 Calibri / 微软雅黑，结构包含封面、目录、执行摘要、章节、
KPI 信息块、财务/估值表、评级总表、60 铲子股行业列表、
图表插入页、附录与免责声明。

用法
----
::

    from morgan_stanley_docx import make_report
    make_report(data, "report.docx", theme="classic", language="zh")
"""

from __future__ import annotations

import os
import re
import tempfile
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm


# =============================================================================
# 1. 主题与颜色系统
# =============================================================================

MS_NAVY = RGBColor(0x0B, 0x2C, 0x5C)       # 品牌主色
MS_GOLD = RGBColor(0xC8, 0xA9, 0x51)       # 辅助强调色
MS_DARK_NAVY = RGBColor(0x06, 0x1A, 0x3A)
OW_GREEN = RGBColor(0x1F, 0x7A, 0x3E)
EW_ORANGE = RGBColor(0xD9, 0x77, 0x06)
UW_RED = RGBColor(0xB9, 0x1C, 0x1C)
TEXT = RGBColor(0x25, 0x1A, 0x1A)            # MS正文色：深棕黑（非纯黑）
MUTED = RGBColor(0x66, 0x66, 0x66)           # 次要文字/标注：中灰色
ALT_BG = RGBColor(0xF5, 0xF6, 0xF8)
LINE = RGBColor(0xD4, 0xD4, 0xD4)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
# ---- MS PDF 学到的颜色 ----
MS_BRAND_BLUE = RGBColor(0x00, 0x55, 0x9F)  # 摩根蓝：标题强调
MS_CHART_BLUE = RGBColor(0x3B, 0x81, 0xB9)  # 图表蓝
MS_CHART_RED = RGBColor(0xFB, 0x03, 0x01)    # 图表红
MS_GREEN_POS = RGBColor(0x00, 0xAF, 0x50)   # 绿色正面
MS_RED_NEG = RGBColor(0xC0, 0x00, 0x00)     # 红色负面
MS_AUX_GRAY = RGBColor(0xAB, 0x9C, 0x99)    # 辅助灰：浅灰棕色
MS_GOLD_LINE = RGBColor(0xC8, 0xA9, 0x51)    # 金色分隔线

# ---- DCF 语义化颜色 ----
DEEP_NAVY = RGBColor(0x1F, 0x38, 0x64)     # 深海军蓝 - 主标题背景
MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)    # 中蓝色 - 副标题/表头背景
INPUT_BG = RGBColor(0xFF, 0xF2, 0xCC)       # 浅黄色 - 输入假设背景
CALC_BG = RGBColor(0xE2, 0xEF, 0xDA)        # 浅绿色 - 关键计算行
SUMMARY_BG = RGBColor(0xD9, 0xE1, 0xF2)     # 浅蓝紫色 - 汇总行
CAPEX_BG = RGBColor(0xFC, 0xE4, 0xD6)      # 浅橙色 - CapEx/ΔNWC行
BEAR_COLOR = RGBColor(0x84, 0x3C, 0x0C)    # 棕橙色 - 熊市标识
BASE_COLOR = RGBColor(0x37, 0x56, 0x23)    # 深绿色 - 基准标识
BULL_COLOR = RGBColor(0x1F, 0x38, 0x64)    # 深海军蓝 - 牛市标识
BEAR_BG = RGBColor(0xFC, 0xE4, 0xD6)       # 浅橙色 - 熊市单元格
BASE_BG = RGBColor(0xE2, 0xEF, 0xDA)        # 浅绿色 - 基准单元格
BULL_BG = RGBColor(0xDD, 0xEE, 0xFF)       # 浅紫色 - 牛市单元格
ALT_ROW = RGBColor(0xF2, 0xF2, 0xF2)        # 浅灰色 - 交替行
HIGHLIGHT_YELLOW = RGBColor(0xFF, 0xFF, 0x00)   # 纯黄色 - 基准交叉点
NEGATIVE_RED = RGBColor(0xC0, 0x00, 0x00)   # 深红色 - 负数
MUTED_GRAY = RGBColor(0x59, 0x59, 0x59)    # 灰色 - 说明文字

THEMES: Dict[str, Dict[str, RGBColor]] = {
    "classic": {
        "navy": MS_NAVY,
        "gold": MS_GOLD,
        "dark_navy": MS_DARK_NAVY,
        "ow": OW_GREEN,
        "ew": EW_ORANGE,
        "uw": UW_RED,
        "text": TEXT,
        "muted": MUTED,
        "alt_bg": ALT_BG,
        "line": LINE,
        "white": WHITE,
        # MS PDF 学到的颜色
        "brand_blue": MS_BRAND_BLUE,
        "chart_blue": MS_CHART_BLUE,
        "chart_red": MS_CHART_RED,
        "green_pos": MS_GREEN_POS,
        "red_neg": MS_RED_NEG,
        "aux_gray": MS_AUX_GRAY,
        "gold_line": MS_GOLD_LINE,
        # DCF 语义化颜色
        "deep_navy": DEEP_NAVY,
        "medium_blue": MEDIUM_BLUE,
        "input_bg": INPUT_BG,
        "calc_bg": CALC_BG,
        "summary_bg": SUMMARY_BG,
        "capex_bg": CAPEX_BG,
        "bear_color": BEAR_COLOR,
        "base_color": BASE_COLOR,
        "bull_color": BULL_COLOR,
        "bear_bg": BEAR_BG,
        "base_bg": BASE_BG,
        "bull_bg": BULL_BG,
        "alt_row": ALT_ROW,
        "highlight_yellow": HIGHLIGHT_YELLOW,
        "negative_red": NEGATIVE_RED,
        "muted_gray": MUTED_GRAY,
    },
}

FONT_EN_TITLE = "Arial"
FONT_EN_BODY = "Arial"
FONT_CN_TITLE = "微软雅黑"
FONT_CN_BODY = "微软雅黑"


def _theme_colors(theme: str) -> Dict[str, RGBColor]:
    return THEMES.get((theme or "").lower(), THEMES["classic"])


# =============================================================================
# 2a. Exhibit 全局计数器
# =============================================================================

_exhibit_counter: int = 0


def _reset_exhibit_counter() -> None:
    """重置 Exhibit 编号计数器（每次生成报告时调用）。"""
    global _exhibit_counter
    _exhibit_counter = 0


def _next_exhibit_label(title: str) -> str:
    """返回 'Exhibit N: title' 格式标签，并递增计数器。"""
    global _exhibit_counter
    _exhibit_counter += 1
    return f"Exhibit {_exhibit_counter}: {title}"


def _exhibit_source(source: Optional[str] = None) -> str:
    """返回 Exhibit 底部来源标注。"""
    return source or "Source: Morgan Stanley Research"


# =============================================================================
# 2. 工具函数 —— XML / 字体 / 单元格 / 段落边框
# =============================================================================

_CN_RE = re.compile(r"[\u4e00-\u9fff\u3400-\u4dbf]")


def _has_cjk(text: str) -> bool:
    return bool(_CN_RE.search(text or ""))


def _hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_run_font(run, size: int = 11, bold: bool = False,
                  italic: bool = False,
                  color: Optional[RGBColor] = None,
                  cjk: bool = False,
                  en_title: bool = False) -> None:
    """对 run 同时设置英文/中文字体。"""
    if cjk:
        run.font.name = FONT_CN_BODY
    elif en_title:
        run.font.name = FONT_EN_TITLE
    else:
        run.font.name = FONT_EN_BODY

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    # 中文 eastAsia，英文 ascii/hAnsi
    rFonts.set(qn("w:eastAsia"), FONT_CN_TITLE if cjk else FONT_CN_BODY)
    rFonts.set(qn("w:ascii"), FONT_EN_TITLE if en_title else FONT_EN_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_EN_TITLE if en_title else FONT_EN_BODY)

    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_run(paragraph, text: str, size: int = 11, bold: bool = False,
             italic: bool = False, color: Optional[RGBColor] = None,
             en_title: bool = False) -> None:
    if text is None:
        return
    run = paragraph.add_run(str(text))
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color,
                  cjk=_has_cjk(str(text)), en_title=en_title)


def _set_cell_shd(cell, color_hex: str) -> None:
    """设置单元格背景填充色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.upper())


def _set_cell_borders(cell, color: RGBColor = LINE,
                      size: str = "4", sides=None) -> None:
    if sides is None:
        sides = ("top", "left", "bottom", "right")
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    hex_c = _hex(color)
    for side in sides:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), hex_c)
        borders.append(node)
    tcPr.append(borders)


def _set_paragraph_border(paragraph, position: str, color: RGBColor,
                          size: str = "8") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    node = OxmlElement(f"w:{position}")
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), size)
    node.set(qn("w:space"), "1")
    node.set(qn("w:color"), _hex(color))
    pbdr.append(node)


def _gold_short_underline(doc, theme_colors: Dict[str, RGBColor],
                          width_inches: float = 1.2) -> None:
    """在当前位置添加一段左侧对齐的短金色装饰线。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # 使用空格占位 + 段落下边框实现短装饰线
    p.paragraph_format.left_indent = Pt(0)
    # 直接在段落内添加一个隐藏字符，并在下边框位置画金色线
    empty_run = p.add_run("\u2003")
    _set_run_font(empty_run, size=1, color=WHITE)
    # 用 tabs + 自定义 border 不行；改为使用一张窄表单元格实现
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.width = Inches(width_inches)
    _set_cell_shd(cell, _hex(theme_colors["gold"]))
    # 移除默认边框，做一条纯填充的短线
    # cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell_para = cell.paragraphs[0]
    cell_para.paragraph_format.space_before = Pt(0)
    cell_para.paragraph_format.space_after = Pt(0)
    _add_run(cell_para, "", size=1)


def _gold_separator_line(doc, theme_colors: Dict[str, RGBColor],
                         thickness_pt: int = 2) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, " ", size=1)
    sz = str(thickness_pt * 8)
    _set_paragraph_border(p, "bottom", theme_colors["gold"], size=sz)


def _navy_separator_line(doc, theme_colors: Dict[str, RGBColor],
                         thickness_pt: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, " ", size=1)
    sz = str(thickness_pt * 8)
    _set_paragraph_border(p, "bottom", theme_colors["navy"], size=sz)


# =============================================================================
# 3. 数字 / 文本格式化
# =============================================================================

def _fmt_number(value: Any) -> str:
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)
    if float(v).is_integer() and not (isinstance(value, str)
                                      and "." in value):
        return f"{int(v):,}"
    return f"{v:,.2f}"


def _fmt_percent(value: Any) -> str:
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{v:.1f}%"


def _fmt_currency(value: Any, currency: str = "USD") -> str:
    try:
        v = float(value)
    except (TypeError, ValueError):
        return str(value)
    symbol = {"USD": "$", "CNY": "¥", "HKD": "HK$", "EUR": "€",
              "GBP": "£"}.get(currency or "USD", "$")
    if float(v).is_integer():
        return f"{symbol}{int(v):,}"
    return f"{symbol}{v:,.2f}"


def _looks_numeric(text: Any) -> bool:
    if not isinstance(text, str):
        return False
    s = text.strip()
    m = re.fullmatch(r"-?\d[\d,]*(\.\d+)?\s*[%x×]?", s)
    if m:
        return True
    m = re.fullmatch(r"[¥$€£]\s*-?\d[\d,]*(\.\d+)?\s*[MB%]?", s)
    return bool(m)


# =============================================================================
# 4.1 图表嵌入（matplotlib）——运行时检测，缺失则退回占位
# =============================================================================

def _try_import_matplotlib():
    """运行时尝试导入 matplotlib。返回 (pyplot_module, ok_flag)。"""
    try:
        import matplotlib  # noqa: F401
        import matplotlib.pyplot as plt  # noqa: F401
        return plt, True
    except Exception:
        return None, False


# ---- Publication-grade chart style constants ----
MS_PALETTE = [
    '#1F3864',  # Deep Navy
    '#2E75B6',  # Medium Blue
    '#C8A951',  # Gold
    '#00AF50',  # Green
    '#E37C2B',  # Orange
    '#B91C1C',  # Red
    '#6B4C9A',  # Purple
    '#4A90D9',  # Light Blue
]

MS_CHART_TITLE_COLOR = '#1F3864'
MS_CHART_SPINE_COLOR = '#4A4A4A'
MS_CHART_TICK_COLOR = '#4A4A4A'
MS_CHART_GRID_COLOR = '#E8E8E8'
MS_CHART_SOURCE_TEXT = 'Source: Morgan Stanley Research'


def _setup_publication_style(plt):
    """Apply publication-grade global style to matplotlib rcParams."""
    plt.rcParams.update({
        'figure.dpi': 300,
        'savefig.dpi': 300,
        'font.family': 'sans-serif',
        'font.sans-serif': [
            'Noto Sans CJK SC', 'Noto Sans CJK',
            'DejaVu Sans', 'Arial',
            'SimHei', 'Microsoft YaHei',
        ],
        'axes.unicode_minus': False,
        'figure.facecolor': '#FFFFFF',
        'axes.facecolor': '#FFFFFF',
        'axes.edgecolor': MS_CHART_SPINE_COLOR,
        'axes.linewidth': 0.6,
        'axes.grid': True,
        'grid.color': MS_CHART_GRID_COLOR,
        'grid.linewidth': 0.5,
        'grid.linestyle': '--',
        'grid.alpha': 1.0,
        'xtick.color': MS_CHART_TICK_COLOR,
        'ytick.color': MS_CHART_TICK_COLOR,
        'xtick.labelsize': 8,
        'ytick.labelsize': 8,
        'axes.spines.top': False,
        'axes.spines.right': False,
        'text.color': '#333333',
    })


def _apply_ax_publication_style(ax, title=''):
    """Apply per-axis publication style tweaks (title, spines, grid)."""
    ax.grid(True, axis='y', color=MS_CHART_GRID_COLOR, linewidth=0.5,
            linestyle='--', zorder=0)
    ax.grid(True, axis='x', color=MS_CHART_GRID_COLOR, linewidth=0.5,
            linestyle='--', zorder=0, alpha=0.3)
    ax.tick_params(colors=MS_CHART_TICK_COLOR, labelsize=8)
    if title:
        ax.set_title(title, fontsize=11, fontweight='bold',
                     color=MS_CHART_TITLE_COLOR, loc='left')


def _add_source_watermark(fig):
    """Add a source footnote at the bottom of the figure."""
    fig.text(0.5, 0.01, MS_CHART_SOURCE_TEXT,
             ha='center', va='bottom', fontsize=7,
             fontstyle='italic', color='#999999')


def _add_bar_value_labels(ax, xs, values, bottoms=None, fmt='{:,.0f}',
                          fontsize=7, color='#4A4A4A', fontweight='bold',
                          offset=0):
    """Add value labels on top of bars."""
    for i, x in enumerate(xs):
        v = values[i]
        b = (bottoms[i] if bottoms else 0) + offset
        if v is not None and v != 0:
            ax.text(x, b + v, fmt.format(v),
                    ha='center', va='bottom', fontsize=fontsize,
                    color=color, fontweight=fontweight, zorder=5)


def _add_mean_reference_line(ax, values, color='#C8A951', linewidth=0.8,
                              linestyle='--', alpha=0.6):
    """Add a horizontal mean reference line."""
    valid = [v for v in values if v is not None]
    if valid:
        mean_val = sum(valid) / len(valid)
        ax.axhline(y=mean_val, color=color, linewidth=linewidth,
                   linestyle=linestyle, alpha=alpha, zorder=1)


def _add_waterfall_connectors(ax, xs, bottoms, heights, color='#999999',
                               linewidth=0.8, linestyle='--'):
    """Add dashed connector lines between waterfall bars."""
    for i in range(len(xs) - 1):
        top_current = bottoms[i] + heights[i]
        bottom_next = bottoms[i + 1]
        ax.plot([xs[i], xs[i + 1]], [top_current, bottom_next],
                color=color, linewidth=linewidth, linestyle=linestyle,
                zorder=1)


def _to_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    # 处理字符串如 "1,210"、"$12"、"8.7%"、"—"
    s = str(value).strip().replace(",", "").replace("%", "")
    s = re.sub(r"[¥$€£]", "", s)
    if s in ("", "—", "-", "NA", "N/A", "nan"):
        return None
    try:
        return float(s)
    except (TypeError, ValueError):
        return None


def _add_chart_image(doc, chart_data: Dict[str, Any],
                     theme_colors: Dict[str, RGBColor],
                     width_inches: float = 6.2) -> str:
    """根据 chart_data 生成一张 PNG 图并嵌入 doc。返回临时文件路径（调用者不负责删除）。

    支持的 chart_data 字段：
        - type: "bar" | "line" | "pie" | "waterfall"（默认 "bar"）
        - title: 标题（可选）
        - 其余字段按类型而定，见函数内注释。

    若 matplotlib 不可用，则插入一个占位段落并返回空字符串。
    """
    if chart_data is None:
        return ""

    plt, ok = _try_import_matplotlib()
    if not ok:
        note = doc.add_paragraph()
        _add_run(note,
                 "[matplotlib 未安装，图表已降级为占位。"
                 "pip install matplotlib 后可渲染]",
                 size=10, italic=True, color=theme_colors["muted"])
        return ""

    ctype = str(chart_data.get("type", "bar")).lower()
    title = str(chart_data.get("title") or "")

    # Apply publication-grade global style
    _setup_publication_style(plt)

    fig, ax = plt.subplots(figsize=(width_inches, 3.6), dpi=300)
    fig.patch.set_facecolor('#FFFFFF')
    ax.set_facecolor('#FFFFFF')

    try:
        if ctype == "bar":
            # Revenue & EBITDA 双柱图：
            # chart_data = {..., "labels": ["FY24","FY25E",...],
            #                "revenue": [4810, 5230, ...],
            #                "ebitda":  [1210, 1405, ...]}
            labels = list(chart_data.get("labels") or [])
            rev = [_to_float(x) for x in (chart_data.get("revenue") or [])]
            ebitda = [_to_float(x) for x in (chart_data.get("ebitda") or [])]
            n = max(len(rev), len(ebitda), len(labels))
            if n == 0:
                raise ValueError("bar chart 缺少数据")
            labels = (labels + [""] * n)[:n]
            rev = (rev + [0.0] * n)[:n]
            ebitda = (ebitda + [0.0] * n)[:n]
            xs = list(range(n))
            w = 0.32
            # Rounded bar effect via white edge
            ax.bar([x - w / 2 for x in xs], rev, width=w, color=MS_PALETTE[0],
                   label="Revenue", edgecolor='white', linewidth=0, zorder=3)
            ax.bar([x + w / 2 for x in xs], ebitda, width=w, color=MS_PALETTE[2],
                   label="EBITDA", edgecolor='white', linewidth=0, zorder=3)
            # Value labels on top of bars
            for i in range(n):
                if rev[i] is not None and rev[i] != 0:
                    ax.text(xs[i] - w / 2, rev[i], f'{rev[i]:,.0f}',
                            ha='center', va='bottom', fontsize=7,
                            color='#4A4A4A', fontweight='bold', zorder=5)
                if ebitda[i] is not None and ebitda[i] != 0:
                    ax.text(xs[i] + w / 2, ebitda[i], f'{ebitda[i]:,.0f}',
                            ha='center', va='bottom', fontsize=7,
                            color='#4A4A4A', fontweight='bold', zorder=5)
            # Mean reference line
            all_vals = [v for v in rev + ebitda if v is not None and v != 0]
            if all_vals:
                mean_val = sum(all_vals) / len(all_vals)
                ax.axhline(y=mean_val, color='#C8A951', linewidth=0.8,
                           linestyle='--', alpha=0.6, zorder=1)
            ax.set_xticks(xs)
            ax.set_xticklabels(labels)
            # Y-axis thousand separator format
            ax.yaxis.set_major_formatter(
                plt.FuncFormatter(lambda x, _: f'{x:,.0f}'))
            ax.legend(frameon=False, fontsize=8, labelcolor=MS_CHART_TICK_COLOR,
                      loc='upper right')

        elif ctype == "line":
            # 利润率折线图：
            # chart_data = {..., "labels": [...],
            #                "gross_margin": [...],
            #                "ebitda_margin": [...],
            #                "net_margin": [...]}
            labels = list(chart_data.get("labels") or [])
            gm = [_to_float(x) for x in (chart_data.get("gross_margin") or [])]
            em = [_to_float(x) for x in (chart_data.get("ebitda_margin") or [])]
            nm = [_to_float(x) for x in (chart_data.get("net_margin") or [])]
            n = max(len(gm), len(em), len(nm), len(labels))
            if n == 0:
                raise ValueError("line chart 缺少数据")
            labels = (labels + [""] * n)[:n]
            gm = (gm + [None] * n)[:n]
            em = (em + [None] * n)[:n]
            nm = (nm + [None] * n)[:n]
            xs = list(range(n))
            # Line styles: linewidth=2.0, markers with different shapes/sizes
            if any(v is not None for v in gm):
                gm_clean = [v if v is not None else float("nan") for v in gm]
                ax.plot(xs, gm_clean, marker="o", markersize=5,
                        color=MS_PALETTE[0], linewidth=2.0,
                        label="Gross Margin", zorder=3)
                # Fill area under first line
                ax.fill_between(xs, 0, gm_clean, alpha=0.08,
                                color=MS_PALETTE[0], zorder=1)
                # Data point labels
                for i, v in enumerate(gm):
                    if v is not None:
                        ax.text(xs[i], v, f'{v:.1f}%', ha='center',
                                va='bottom', fontsize=7,
                                color=MS_PALETTE[0], fontweight='bold', zorder=5)
            if any(v is not None for v in em):
                em_clean = [v if v is not None else float("nan") for v in em]
                ax.plot(xs, em_clean, marker="s", markersize=4,
                        color=MS_PALETTE[2], linewidth=2.0,
                        label="EBITDA Margin", zorder=3)
                for i, v in enumerate(em):
                    if v is not None:
                        ax.text(xs[i], v, f'{v:.1f}%', ha='center',
                                va='bottom', fontsize=7,
                                color=MS_PALETTE[2], fontweight='bold', zorder=5)
            if any(v is not None for v in nm):
                nm_clean = [v if v is not None else float("nan") for v in nm]
                ax.plot(xs, nm_clean, marker="^", markersize=5,
                        color=MS_PALETTE[3], linewidth=2.0,
                        label="Net Margin", zorder=3)
                for i, v in enumerate(nm):
                    if v is not None:
                        ax.text(xs[i], v, f'{v:.1f}%', ha='center',
                                va='bottom', fontsize=7,
                                color=MS_PALETTE[3], fontweight='bold', zorder=5)
            ax.set_xticks(xs)
            ax.set_xticklabels(labels)
            ax.set_ylim(0, 100)
            ax.yaxis.set_major_locator(plt.MultipleLocator(10))
            ax.legend(frameon=False, fontsize=8, labelcolor=MS_CHART_TICK_COLOR,
                      loc='upper right')
            ax.set_ylabel("%", fontsize=9, color=MS_CHART_TICK_COLOR)

        elif ctype == "pie":
            # 行业配置饼图：
            # chart_data = {..., "labels": [...], "values": [...]}
            labels = list(chart_data.get("labels") or [])
            values = [_to_float(x) for x in (chart_data.get("values") or [])]
            values = [v for v in values if v is not None]
            if not values or len(values) != len(labels):
                raise ValueError("pie chart 缺少数据")
            colors = [MS_PALETTE[i % len(MS_PALETTE)] for i in range(len(values))]
            # Explode the largest slice
            max_idx = values.index(max(values))
            explodes = [0.05 if i == max_idx else 0 for i in range(len(values))]
            wedges, texts, autotexts = ax.pie(
                values, labels=None, colors=colors,
                autopct='%1.1f%%', startangle=90,
                pctdistance=0.75,
                explode=explodes,
                shadow=True,
                wedgeprops={'linewidth': 1.5, 'edgecolor': 'white'},
                textprops={'fontsize': 9, 'color': '#1A1A1A'})
            # Bold percentage text
            for at in autotexts:
                at.set_fontsize(9)
                at.set_fontweight('bold')
            # Legend on the right
            ax.legend(wedges, labels, title="Segments",
                      loc="center left", bbox_to_anchor=(1, 0, 0.5, 1),
                      fontsize=8, frameon=False,
                      title_fontsize=8)
            ax.set_aspect("equal")

        elif ctype == "waterfall":
            # 估值桥梁图（当前价 → 目标价）：
            # chart_data = {..., "labels": [...], "values": [...],
            #                "current": 98.5, "target": 125.0}
            labels = list(chart_data.get("labels") or [])
            values = [_to_float(x) for x in (chart_data.get("values") or [])]
            current = _to_float(chart_data.get("current")) or 0.0
            target = _to_float(chart_data.get("target")) or 0.0
            bars = ["Current"] + [str(l) for l in labels] + ["Target"]
            # 起始柱 = current；中间 delta；终点柱 = target
            bar_values = [current] + values + [target]
            # 以堆叠方式渲染（累计）
            cumulative = 0.0
            bottoms = []
            heights = []
            for i, v in enumerate(bar_values):
                if i == 0:
                    bottoms.append(0.0)
                    heights.append(float(v or 0.0))
                    cumulative = float(v or 0.0)
                elif i == len(bar_values) - 1:
                    bottoms.append(0.0)
                    heights.append(float(v or 0.0))
                else:
                    bottoms.append(cumulative)
                    heights.append(float(v or 0.0))
                    cumulative += float(v or 0.0)
            colors_bar = []
            for i, h in enumerate(heights):
                if i == 0:
                    colors_bar.append(MS_PALETTE[0])  # Deep Navy
                elif i == len(heights) - 1:
                    colors_bar.append(MS_PALETTE[2])  # Gold
                else:
                    colors_bar.append(MS_PALETTE[3] if h >= 0 else MS_PALETTE[5])  # Green/Red
            xs = list(range(len(bars)))
            ax.bar(xs, heights, bottom=bottoms, width=0.5, color=colors_bar,
                   edgecolor='white', linewidth=0, zorder=3)
            # Connector lines between bars
            _add_waterfall_connectors(ax, xs, bottoms, heights)
            ax.set_xticks(xs)
            ax.set_xticklabels(bars, rotation=20, ha="right", fontsize=8)
            # Value labels: positive=green, negative=red
            for i, (x, h, b) in enumerate(zip(xs, heights, bottoms)):
                value = h if i in (0, len(heights) - 1) else h
                sign = "+" if value > 0 and i not in (0, len(heights) - 1) else ""
                lbl_color = MS_PALETTE[3] if value > 0 and i not in (0, len(heights) - 1) else (
                    MS_PALETTE[5] if value < 0 and i not in (0, len(heights) - 1) else '#4A4A4A')
                ax.text(x, b + h, f"{sign}{value:.1f}",
                        ha="center", va="bottom", fontsize=8,
                        color=lbl_color, fontweight='bold', zorder=5)

        else:
            # 不支持的类型 -> 占位文字
            plt.close(fig)
            note = doc.add_paragraph()
            _add_run(note, f"[chart type='{ctype}' 未支持]",
                     size=10, italic=True, color=theme_colors["muted"])
            return ""

        # Apply publication title style
        _apply_ax_publication_style(ax, title=title)
        # Source watermark
        _add_source_watermark(fig)
        fig.tight_layout(pad=1.5)

        # 写入临时文件
        fd, path = tempfile.mkstemp(prefix="ms_chart_", suffix=".png")
        os.close(fd)
        try:
            fig.savefig(path, dpi=300, facecolor='#FFFFFF',
                        bbox_inches='tight')
        finally:
            plt.close(fig)

        try:
            doc.add_picture(path, width=Inches(width_inches))
        except Exception:
            # docx 嵌入失败
            note = doc.add_paragraph()
            _add_run(note, "[图表嵌入失败]", size=10, italic=True,
                     color=theme_colors["muted"])
            try:
                os.remove(path)
            except OSError:
                pass
            return ""

        return path  # 保留文件供调用者自行决定是否删除
    except Exception as exc:
        plt.close(fig)
        note = doc.add_paragraph()
        _add_run(note, f"[图表渲染失败: {exc}]", size=10, italic=True,
                 color=theme_colors["muted"])
        return ""


def _cleanup_chart_paths(paths: List[str]) -> None:
    """删除临时图表文件，忽略错误。"""
    for p in paths:
        if not p:
            continue
        try:
            os.remove(p)
        except OSError:
            pass


# =============================================================================
# 5. Cover 封面页（MS 标准模板）
# =============================================================================

def _cover_page(doc, data: Dict[str, Any],
                theme_colors: Dict[str, RGBColor]) -> None:
    """MS 标准封面页模板。

    布局：
      顶部: MS Logo 区域（深蓝色矩形 + 白色 MS 文字）
      报告类型标签 + 日期 + 行业/地区 + 法人实体
      分析师信息区
      金色分隔线
      主标题（大字号，深棕黑 #251A1A）
      副标题/核心论点
      金色分隔线
      Key Takeaways（bullet 列表）
      What's Changed 表格（可选）
      行业观点标签（可选）
      底部免责声明（小字灰色）
    """
    brand_blue = theme_colors.get("brand_blue", MS_BRAND_BLUE)
    gold = theme_colors.get("gold_line", MS_GOLD_LINE)
    text_color = theme_colors["text"]
    muted = theme_colors["muted"]
    aux_gray = theme_colors.get("aux_gray", MS_AUX_GRAY)

    # ---- 1. 顶部 MS Logo 区域 ----
    # 深蓝色矩形 + 白色 "MS" 文字模拟 Logo
    logo_block = doc.add_table(rows=1, cols=1)
    logo_block.alignment = WD_TABLE_ALIGNMENT.LEFT
    logo_cell = logo_block.rows[0].cells[0]
    logo_cell.width = Inches(6.5)
    logo_cell.height = Inches(0.55)
    _set_cell_shd(logo_cell, _hex(brand_blue))
    _set_cell_borders(logo_cell, color=brand_blue, size="0",
                      sides=("top", "left", "right"))
    _set_cell_borders(logo_cell, color=gold, size="6",
                      sides=("bottom",))

    logo_para = logo_cell.paragraphs[0]
    logo_para.paragraph_format.space_before = Pt(6)
    logo_para.paragraph_format.space_after = Pt(6)
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 左侧：白色 "MS" + Morgan Stanley
    _add_run(logo_para, "MS", size=22, bold=True,
             color=WHITE, en_title=True)
    _add_run(logo_para, "  MORGAN STANLEY", size=11, bold=True,
             color=WHITE, en_title=True)

    # 右侧报告类型标签
    cat_para = logo_cell.add_paragraph()
    cat_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cat_para.paragraph_format.space_before = Pt(0)
    cat_para.paragraph_format.space_after = Pt(0)
    research_type = data.get("research_type") or "Foundation Research"
    _add_run(cat_para, str(research_type).upper(), size=10, bold=True,
             color=WHITE, en_title=True)

    # ---- 2. 日期 | 行业/地区 | 法人实体 ----
    meta_line = doc.add_paragraph()
    meta_line.paragraph_format.space_before = Pt(8)
    meta_line.paragraph_format.space_after = Pt(2)
    date_str = data.get("date_str") or ""
    industry = data.get("industry") or data.get("sector") or ""
    entity = data.get("entity") or "Morgan Stanley & Co."
    parts = [p for p in [date_str, industry, entity] if p]
    _add_run(meta_line, "  |  ".join(parts), size=9,
             color=muted, en_title=True)

    # ---- 3. 分析师信息区 ----
    analyst = data.get("analyst") or ""
    analyst_title = data.get("analyst_title") or "Equity Research Analyst"
    analyst_email = data.get("analyst_email") or data.get("contact") or ""
    if analyst:
        ana_p = doc.add_paragraph()
        ana_p.paragraph_format.space_before = Pt(4)
        ana_p.paragraph_format.space_after = Pt(2)
        _add_run(ana_p, analyst, size=11, bold=True,
                 color=text_color)
        if analyst_title:
            _add_run(ana_p, f"  ({analyst_title})", size=10,
                     color=muted)
        if analyst_email:
            _add_run(ana_p, f"  |  {analyst_email}", size=10,
                     color=muted)

    # ---- 4. 金色分隔线 ----
    _gold_separator_line(doc, theme_colors, thickness_pt=2)

    # ---- 5. 主标题 ----
    title_cn = data.get("title_cn") or ""
    title_en = data.get("title_en") or ""
    if title_cn:
        t1 = doc.add_paragraph()
        t1.paragraph_format.space_before = Pt(8)
        t1.paragraph_format.space_after = Pt(4)
        _add_run(t1, title_cn, size=26, bold=True, color=text_color)
    if title_en:
        t2 = doc.add_paragraph()
        t2.paragraph_format.space_before = Pt(0)
        t2.paragraph_format.space_after = Pt(4)
        _add_run(t2, title_en, size=16, bold=False,
                 color=text_color, en_title=True)

    # ---- 6. 副标题/核心论点 ----
    subtitle = data.get("subtitle") or ""
    if subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(6)
        _add_run(sub, subtitle, size=12, italic=True, color=muted)

    # ---- 7. 金色分隔线 ----
    _gold_separator_line(doc, theme_colors, thickness_pt=2)

    # ---- 8. Key Takeaways ----
    key_takeaways = data.get("key_takeaways") or []
    if key_takeaways:
        kt_head = doc.add_paragraph()
        kt_head.paragraph_format.space_before = Pt(6)
        kt_head.paragraph_format.space_after = Pt(4)
        _add_run(kt_head, "Key Takeaways", size=13, bold=True,
                 color=brand_blue)

        for item in key_takeaways:
            bullet = doc.add_paragraph()
            bullet.paragraph_format.space_before = Pt(1)
            bullet.paragraph_format.space_after = Pt(1)
            bullet.paragraph_format.left_indent = Pt(18)
            _add_run(bullet, "\u2022  ", size=10, bold=True,
                     color=brand_blue)
            _add_run(bullet, str(item), size=10, color=text_color)

    # ---- 9. What's Changed 表格 ----
    whats_changed = data.get("whats_changed") or []
    if whats_changed:
        doc.add_paragraph()  # 空行
        wc_head = doc.add_paragraph()
        wc_head.paragraph_format.space_before = Pt(4)
        wc_head.paragraph_format.space_after = Pt(4)
        _add_run(wc_head, "What's Changed", size=13, bold=True,
                 color=brand_blue)

        wc_table = doc.add_table(rows=1 + len(whats_changed), cols=3)
        wc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        wc_headers = ["Item", "From", "To"]
        for c_idx, h in enumerate(wc_headers):
            tc = wc_table.rows[0].cells[c_idx]
            _set_cell_shd(tc, _hex(brand_blue))
            _add_run(tc.paragraphs[0], h, size=9, bold=True,
                     color=WHITE, en_title=True)

        for r_idx, row_data in enumerate(whats_changed):
            row_idx = r_idx + 1
            if isinstance(row_data, dict):
                vals = [
                    str(row_data.get("item") or ""),
                    str(row_data.get("from_val") or row_data.get("from") or ""),
                    str(row_data.get("to_val") or row_data.get("to") or ""),
                ]
            elif isinstance(row_data, (list, tuple)):
                vals = [str(v) for v in row_data[:3]]
            else:
                vals = [str(row_data), "", ""]
            for c_idx, val in enumerate(vals):
                tc = wc_table.rows[row_idx].cells[c_idx]
                if r_idx % 2 == 1:
                    _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
                _add_run(tc.paragraphs[0], val, size=9,
                         color=text_color)

    # ---- 10. 行业观点标签 ----
    industry_view = data.get("industry_view") or ""
    if industry_view:
        doc.add_paragraph()
        iv_p = doc.add_paragraph()
        iv_p.paragraph_format.space_before = Pt(4)
        iv_p.paragraph_format.space_after = Pt(4)
        iv_upper = str(industry_view).strip().upper()
        # 颜色映射
        if iv_upper in ("ATTRACTIVE",):
            iv_color = theme_colors.get("green_pos", MS_GREEN_POS)
        elif iv_upper in ("IN-LINE", "IN-LINE", "IN LINE"):
            iv_color = theme_colors.get("ew", EW_ORANGE)
        elif iv_upper in ("CAUTIOUS",):
            iv_color = theme_colors.get("red_neg", MS_RED_NEG)
        else:
            iv_color = brand_blue
        _add_run(iv_p, "Industry View:  ", size=11, bold=True,
                 color=muted)
        _add_run(iv_p, iv_upper, size=14, bold=True, color=iv_color)

    # ---- 11. 评级/目标价（如果有） ----
    rating = data.get("rating")
    target_price = data.get("target_price")
    current_price = data.get("current_price")
    if rating or target_price:
        doc.add_paragraph()
        if rating:
            rate_line = doc.add_paragraph()
            rate_line.paragraph_format.space_after = Pt(2)
            _add_run(rate_line, "Rating  ", size=10, color=muted)
            _add_run(rate_line, str(rating).upper(), size=16,
                     bold=True, color=brand_blue)
        if target_price is not None:
            tp_line = doc.add_paragraph()
            tp_line.paragraph_format.space_after = Pt(2)
            _add_run(tp_line, "Target Price  ", size=10, color=muted)
            _add_run(tp_line,
                     _fmt_currency(target_price,
                                   data.get("currency", "USD")),
                     size=20, bold=True, color=gold)
        if current_price is not None:
            cp_line = doc.add_paragraph()
            cp_line.paragraph_format.space_after = Pt(2)
            _add_run(cp_line, "Last Close  ", size=10, color=muted)
            _add_run(cp_line,
                     _fmt_currency(current_price,
                                   data.get("currency", "USD")),
                     size=14, bold=True, color=text_color)

    # ---- 12. 底部免责声明（小字灰色） ----
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.paragraph_format.space_before = Pt(6)
    disc.paragraph_format.space_after = Pt(2)
    disc_text = (
        "This report has been prepared by Morgan Stanley Research. "
        "It is not intended for distribution to, or use by, any person "
        "or entity in any jurisdiction where such distribution or use would "
        "be contrary to law or regulation."
    )
    _add_run(disc, disc_text, size=7, italic=True, color=aux_gray)

    doc.add_page_break()


# =============================================================================
# 5. Table of Contents 目录页
# =============================================================================

def _toc_page(doc, data: Dict[str, Any],
              theme_colors: Dict[str, RGBColor]) -> None:
    head = doc.add_paragraph(style='Heading 1')
    head.paragraph_format.space_before = Pt(6)
    _add_run(head, "目录 / Table of Contents", size=20, bold=True,
             color=theme_colors["brand_blue"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    # 手动列出预期后续章节（Word 的 TOC 域在打开时更新，此处提供静态版本）
    chapters: List[Tuple[str, str]] = []

    if data.get("executive_summary") or data.get("key_takeaways"):
        chapters.append(("执行摘要 / Executive Summary", "—"))
    if data.get("thesis_charts"):
        n_tc = len(data["thesis_charts"])
        chapters.append((f"Our Thesis in {n_tc} Charts", "—"))
    for section in (data.get("sections") or []):
        t_cn = section.get("title_cn") or "章节"
        t_en = section.get("title_en") or ""
        label = f"{t_cn} / {t_en}" if t_en else t_cn
        chapters.append((label, "—"))
    if data.get("metrics"):
        chapters.append(("关键指标 / Key KPIs", "—"))
    if data.get("financial_table"):
        chapters.append(("财务与估值 / Financials & Valuation", "—"))
    if data.get("rating_table"):
        chapters.append(("评级总表 / Rating Summary", "—"))
    if data.get("shovel_stocks"):
        chapters.append(("60 铲子股行业列表 / Shovel Stocks Universe", "—"))
    if data.get("value_chain"):
        chapters.append(("产业链分析 / Value Chain Analysis", "—"))
    if data.get("shovel_stocks_list"):
        chapters.append(("概念股列表 / Shovel Stocks List", "—"))
    if data.get("scenarios"):
        chapters.append(("情景分析 / Scenario Analysis", "—"))
        chapters.append(("WACC 分析 / WACC Analysis", "—"))
    if data.get("sensitivity"):
        chapters.append(("敏感性分析 / Sensitivity Analysis", "—"))
    if data.get("comparable_companies"):
        chapters.append(("可比公司分析 / Comparable Companies", "—"))
    if data.get("scenarios"):
        chapters.append(("估值桥 / Valuation Bridge", "—"))
    if data.get("exhibits"):
        chapters.append(("Exhibits", "—"))
    chapters.append(("附录 / Appendix", "—"))
    chapters.append(("免责声明 / Disclosure", "—"))

    for idx, (name, _page) in enumerate(chapters, start=1):
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(4)
        _add_run(line, f"{idx:02d}. ", size=12, bold=True,
                 color=theme_colors["gold"])
        _add_run(line, name, size=12, color=theme_colors["navy"])

    doc.add_page_break()


# =============================================================================
# 6. Executive Summary / 执行摘要
# =============================================================================

def _normalize_bullets(value: Any) -> List[str]:
    """将各种输入（字符串按换行/列表）统一为 bullet 字符串列表。"""
    if value is None:
        return []
    if isinstance(value, (list, tuple)):
        return [str(v).strip() for v in value if str(v).strip()]
    text = str(value).strip()
    if not text:
        return []
    # 支持按单/双换行分段
    parts = re.split(r"\n\s*\n|\n", text)
    return [p.strip() for p in parts if p.strip()]


def _exec_summary_page(doc, data: Dict[str, Any],
                       theme_colors: Dict[str, RGBColor]) -> None:
    """执行摘要页：优先读取 data.executive_summary 四象限字典。

    data["executive_summary"] 可为 dict（新风格四象限），或 string（旧风格纯文本）。
    旧风格下自动回退到段落模式；并补充 key_takeaways 的要点列表。
    """
    es = data.get("executive_summary")
    takeaways = data.get("key_takeaways") or []

    # 四象限新风格：必须是 dict 且包含至少一个象限字段
    if isinstance(es, dict) and any(k in es for k in
                                    ("thesis", "tp_and_upside",
                                     "risks", "catalysts")):
        thesis = _normalize_bullets(es.get("thesis"))
        tp_and_upside = _normalize_bullets(es.get("tp_and_upside"))
        risks = _normalize_bullets(es.get("risks"))
        catalysts = _normalize_bullets(es.get("catalysts"))

        # 标题行
        head = doc.add_paragraph(style='Heading 1')
        head.paragraph_format.space_before = Pt(6)
        _add_run(head, "执行摘要 / Executive Summary", size=20, bold=True,
                 color=theme_colors["navy"])
        _gold_short_underline(doc, theme_colors)
        doc.add_paragraph()

        # 2x2 无边框嵌套表实现四象限布局
        outer = doc.add_table(rows=2, cols=2)
        outer.alignment = WD_TABLE_ALIGNMENT.LEFT
        # 去掉所有边框，避免与页面干扰
        for row in outer.rows:
            for cell in row.cells:
                _set_cell_borders(cell, color=WHITE, size="0")

        quadrants = [
            # (row, col, 金色标题, bullets)
            (0, 0, "核心论点 / Thesis", thesis),
            (0, 1, "目标价 & 上行 / TP & Upside", tp_and_upside),
            (1, 0, "关键风险 / Key Risks", risks),
            (1, 1, "催化剂 / Catalysts", catalysts),
        ]
        for r, c, title_cn, bullets in quadrants:
            cell = outer.rows[r].cells[c]
            # 标题
            title_para = cell.paragraphs[0]
            title_para.paragraph_format.space_after = Pt(4)
            _add_run(title_para, title_cn, size=12, bold=True,
                     color=theme_colors["gold"])
            # 内容 bullets
            if not bullets:
                empty_para = cell.add_paragraph()
                _add_run(empty_para, "—", size=10,
                         color=theme_colors["muted"])
                continue
            for b in bullets:
                line = cell.add_paragraph()
                line.paragraph_format.space_after = Pt(2)
                # 统一使用小圆点前缀
                cleaned = b
                if cleaned.startswith(("- ", "• ", "* ")):
                    cleaned = cleaned[2:]
                _add_run(line, "• ", size=10, bold=True,
                         color=theme_colors["gold"])
                _add_run(line, cleaned, size=10,
                         color=theme_colors["text"])

        # 情景概览条（在四象限下方，如果 data 含 scenarios）
        scenarios = data.get("scenarios")
        if scenarios and isinstance(scenarios, dict):
            bear = scenarios.get("bear")
            base = scenarios.get("base")
            bull = scenarios.get("bull")
            if bear and base and bull:
                doc.add_paragraph()
                # 情景概览标题
                sc_head = doc.add_paragraph()
                sc_head.paragraph_format.space_after = Pt(4)
                _add_run(sc_head, "情景概览 / Scenario Overview", size=12,
                         bold=True, color=theme_colors["navy"])
                _gold_short_underline(doc, theme_colors, width_inches=0.8)

                # 三行情景对比表
                currency = data.get("currency", "USD")
                sym = {"USD": "$", "CNY": "¥", "HKD": "HK$",
                       "EUR": "€", "GBP": "£"}.get(currency, "$")
                sc_table = doc.add_table(rows=3, cols=3)
                sc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                sc_table.autofit = False

                scenario_rows = [
                    (bear, "熊市 / Bear", theme_colors["uw"]),
                    (base, "基准 / Base", theme_colors["gold"]),
                    (bull, "牛市 / Bull", theme_colors["ow"]),
                ]
                for r_idx, (s_data, label, accent_color) in enumerate(scenario_rows):
                    # 情景标签列
                    tc0 = sc_table.rows[r_idx].cells[0]
                    tc0.width = Inches(1.5)
                    if r_idx == 1:
                        _set_cell_shd(tc0, "FFF8E7")
                    _add_run(tc0.paragraphs[0], label, size=10, bold=True,
                             color=accent_color)

                    # Per Share Value 列
                    tc1 = sc_table.rows[r_idx].cells[1]
                    tc1.width = Inches(2.0)
                    if r_idx == 1:
                        _set_cell_shd(tc1, "FFF8E7")
                    ps = s_data.get("per_share_value")
                    ps_str = _fmt_currency(ps, currency) if ps else "—"
                    _add_run(tc1.paragraphs[0], ps_str, size=14, bold=True,
                             color=accent_color)

                    # EV/EBITDA 列
                    tc2 = sc_table.rows[r_idx].cells[2]
                    tc2.width = Inches(2.0)
                    if r_idx == 1:
                        _set_cell_shd(tc2, "FFF8E7")
                    ev_ebitda = s_data.get("ev_ebitda")
                    ev_ebitda_str = f"{ev_ebitda}x" if ev_ebitda else "—"
                    _add_run(tc2.paragraphs[0], f"EV/EBITDA: {ev_ebitda_str}",
                             size=10, color=theme_colors["text"])

                # 基准行金色边框
                for c_idx in range(3):
                    _set_cell_borders(sc_table.rows[1].cells[c_idx],
                                      color=theme_colors["gold"], size="8",
                                      sides=("top", "bottom"))

        doc.add_page_break()
        return

    # 兼容旧逻辑：executive_summary（字符串）+ key_takeaways（列表）
    if not es and not takeaways:
        return

    head = doc.add_paragraph(style='Heading 1')
    head.paragraph_format.space_before = Pt(6)
    _add_run(head, "执行摘要 / Executive Summary", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    if isinstance(es, str):
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", es)
                      if p.strip()]
        for i, para in enumerate(paragraphs):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            pPr = p._p.get_or_add_pPr()
            ind = pPr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                pPr.append(ind)
            ind.set(qn("w:firstLineChars"), "200")

            if i == 0 and para:
                first_char = para[0]
                rest = para[1:]
                drop = p.add_run(first_char)
                _set_run_font(drop, size=28, bold=True,
                              color=theme_colors["gold"],
                              cjk=_has_cjk(first_char))
                _add_run(p, rest, size=11, color=theme_colors["text"])
            else:
                _add_run(p, para, size=11, color=theme_colors["text"])

    if takeaways:
        doc.add_paragraph()
        sub = doc.add_paragraph()
        _add_run(sub, "核心要点 / Key Takeaways", size=14, bold=True,
                 color=theme_colors["navy"])
        _gold_separator_line(doc, theme_colors, thickness_pt=1)
        for idx, item in enumerate(takeaways, start=1):
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(4)
            _add_run(line, f"{idx:02d}  ", size=12, bold=True,
                     color=theme_colors["gold"])
            _add_run(line, str(item), size=11, color=theme_colors["text"])

    doc.add_page_break()


# 保持向后兼容：旧函数名作为别名
_executive_summary = _exec_summary_page


# =============================================================================
# 7. Section 章节页
# =============================================================================

def _detect_bullet(text: str) -> Tuple[int, str]:
    """识别双层 bullet。返回 (level, 剩余文本)。
    level=0 代表普通段落；level=1 代表一级 bullet；level=2 代表二级 bullet。
    """
    if not isinstance(text, str):
        return 0, text or ""
    stripped = text.lstrip()
    # 二级：以 "  - " / "  • " 开头（前面至少一个空格或 tab）
    if re.match(r"^(?:\s{2,}|\t)\s*(?:-|•|\*)\s+", text):
        m = re.match(r"^(?:\s+)(?:-|•|\*)\s+", text)
        return 2, text[m.end():] if m else text
    if stripped.startswith(("- ", "• ", "* ")):
        return 1, stripped[2:]
    return 0, text


def _body_paragraph(doc, text: str,
                    theme_colors: Dict[str, RGBColor],
                    lead_word: Optional[str] = None) -> None:
    text = str(text or "")
    level, content = _detect_bullet(text)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25

    if level > 0:
        # Bullet 样式：左侧缩进；使用金色小圆点或横杠作为项目符号
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        if level == 1:
            ind.set(qn("w:leftChars"), "200")
            ind.set(qn("w:firstLineChars"), "-200")
        else:
            ind.set(qn("w:leftChars"), "400")
            ind.set(qn("w:firstLineChars"), "-200")

        bullet_char = "•" if level == 1 else "–"
        _add_run(p, f"{bullet_char} ", size=11, bold=True,
                 color=theme_colors["gold"])
        _add_run(p, content, size=11, color=theme_colors["text"])
        return

    # 普通段落：中文首行缩进 2 字符
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), "200")

    if lead_word:
        _add_run(p, f"{lead_word} — ", size=11, bold=True,
                 color=theme_colors["navy"])
    _add_run(p, text, size=11, color=theme_colors["text"])


def _section(doc, section: Dict[str, Any],
             theme_colors: Dict[str, RGBColor],
             section_index: int = 0,
             sectors_allocation: Optional[Dict[str, Any]] = None) -> None:
    title_cn = section.get("title_cn") or "章节"
    title_en = section.get("title_en") or ""

    # H1 中文 + 英文副标题
    h1 = doc.add_paragraph(style='Heading 1')
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(2)
    _add_run(h1, title_cn, size=20, bold=True, color=theme_colors["navy"])

    if title_en:
        h1_en = doc.add_paragraph()
        h1_en.paragraph_format.space_after = Pt(4)
        _add_run(h1_en, title_en, size=13, italic=True,
                 color=theme_colors["navy"], en_title=True)

    # 段落前金色装饰线
    _gold_short_underline(doc, theme_colors)

    # 章节正文段落（可能含双层 bullet）
    lead_words = ["We believe", "We estimate", "Crucially",
                  "Importantly", "Looking ahead"]
    for i, para in enumerate(section.get("paragraphs") or []):
        lead = lead_words[i % len(lead_words)] if i < len(lead_words) else None
        _body_paragraph(doc, str(para), theme_colors, lead_word=lead)

    # 子章节 / 二级标题 / 图表 / 表格
    for sub in (section.get("subsections") or []):
        if sub.get("title"):
            h2 = doc.add_paragraph(style='Heading 2')
            h2.paragraph_format.space_before = Pt(10)
            h2.paragraph_format.space_after = Pt(4)
            _add_run(h2, sub["title"], size=16, bold=True,
                     color=theme_colors["navy"])
            _gold_short_underline(doc, theme_colors, width_inches=0.8)
        for para in (sub.get("paragraphs") or []):
            _body_paragraph(doc, str(para), theme_colors)

        if sub.get("table"):
            rows_data = sub["table"]
            if rows_data:
                _simple_table(doc, rows_data, theme_colors)

        if sub.get("chart_img"):
            try:
                doc.add_picture(sub["chart_img"], width=Inches(6.0))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(caption,
                         sub.get("chart_caption") or "Exhibit: Chart",
                         size=10, italic=True, color=theme_colors["muted"])
            except Exception:
                pass

    # 在第三章自动插入饼图 exhibit（当 data 含 sectors_allocation 时）
    if sectors_allocation and section_index == 2:
        doc.add_paragraph()
        ex_head = doc.add_paragraph()
        _add_run(ex_head,
                 str(sectors_allocation.get("title")
                     or "行业配置 / Sector Allocation"),
                 size=16, bold=True, color=theme_colors["navy"])
        _gold_short_underline(doc, theme_colors)
        doc.add_paragraph()

        labels = list(sectors_allocation.get("labels") or [])
        values_raw = list(sectors_allocation.get("values") or [])
        # 允许 values 为纯数字；若同时存在 percentages 也可
        values = [_to_float(v) for v in values_raw]
        if any(v is not None for v in values) and len(values) == len(labels):
            chart_data = {
                "type": "pie",
                "title": sectors_allocation.get("title")
                         or "Sector Allocation",
                "labels": labels,
                "values": values,
            }
            _add_chart_image(doc, chart_data, theme_colors)
        else:
            note = doc.add_paragraph()
            _add_run(note, "[无法渲染饼图：labels 与 values 不匹配]",
                     size=10, italic=True, color=theme_colors["muted"])

    doc.add_paragraph()


# =============================================================================
# 8. 简单表格（章节内 / 附录内）
# =============================================================================

def _simple_table(doc, rows: List[List[Any]],
                  theme_colors: Dict[str, RGBColor],
                  header_fill: Optional[str] = None,
                  alt: bool = True) -> None:
    if not rows:
        return
    header_fill = header_fill or _hex(theme_colors["navy"])
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = row[c_idx] if c_idx < len(row) else ""
            tc = table.rows[r_idx].cells[c_idx]
            tc_para = tc.paragraphs[0]
            if r_idx == 0:
                _set_cell_shd(tc, header_fill)
                _add_run(tc_para, str(cell), size=10, bold=True,
                         color=WHITE)
            else:
                if alt and (r_idx % 2 == 0):
                    _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
                _add_run(tc_para, str(cell), size=10,
                         color=theme_colors["text"])


# =============================================================================
# 9. KPI Blocks 关键数字信息块
# =============================================================================

def _kpi_blocks(doc, data: Dict[str, Any],
                theme_colors: Dict[str, RGBColor]) -> None:
    metrics = data.get("metrics") or []
    if not metrics:
        return

    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "关键指标 / Key KPIs", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    # 每行最多 3 个 KPI 卡片
    per_row = 3
    total = len(metrics)
    rows = (total + per_row - 1) // per_row

    for r in range(rows):
        row_items = metrics[r * per_row:(r + 1) * per_row]
        table = doc.add_table(rows=1, cols=per_row)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        for i, cell in enumerate(table.rows[0].cells):
            cell.width = Inches(6.5 / per_row)
            if i >= len(row_items):
                _add_run(cell.paragraphs[0], "", size=1)
                continue
            m = row_items[i]
            _set_cell_shd(cell, _hex(theme_colors["alt_bg"]))
            _set_cell_borders(cell, color=theme_colors["gold"], size="4")

            value = str(m.get("value", m.get("metric", "")))
            unit = str(m.get("unit") or m.get("label_en") or "")
            label = str(m.get("label") or m.get("label_cn") or "")

            p1 = cell.paragraphs[0]
            _add_run(p1, value, size=32, bold=True,
                     color=theme_colors["gold"])
            p2 = cell.add_paragraph()
            _add_run(p2, unit, size=11, bold=True,
                     color=theme_colors["navy"])
            p3 = cell.add_paragraph()
            _add_run(p3, label, size=9, color=theme_colors["muted"])
        doc.add_paragraph()

    doc.add_page_break()


# =============================================================================
# 10. Financial Table 财务/估值表格
# =============================================================================

def _row_key(row: List[Any]) -> str:
    """从行首列提取关键字，用于匹配行类型。"""
    if not row:
        return ""
    key = str(row[0])
    key = re.sub(r"[\s\-—_,()（）$¥€£%]", "", key)
    return key.lower()


def _is_margin_row(row: List[Any]) -> bool:
    key = _row_key(row)
    return any(kw in key for kw in ("毛利率", "grossmargin",
                                    "营业利润率", "operatingmargin",
                                    "ebitdamargin",
                                    "净利润率", "netmargin",
                                    "margin"))


def _is_growth_row(row: List[Any]) -> bool:
    key = _row_key(row)
    return any(kw in key for kw in ("同比", "yoy", "增长"))


def _is_capex_row(row: List[Any]) -> bool:
    key = _row_key(row)
    return "capex" in key or "资本支出" in key


def _is_revenue_row(row: List[Any]) -> bool:
    key = _row_key(row)
    return any(kw in key for kw in ("营业收入", "营业总收入",
                                    "销售收入", "revenue",
                                    "sales", "netrevenue"))


def _fmt_yoy(value: Optional[float]) -> str:
    if value is None:
        return "—"
    return f"{value:+.1f}%"


def _auto_yoy_from_rows(rows: List[List[Any]]) -> Optional[List[Any]]:
    """从行中查找营业收入行，基于其数值生成 YoY 列。"""
    rev_row = None
    for r in rows:
        if _is_revenue_row(r):
            rev_row = r
            break
    if rev_row is None:
        return None
    values: List[Optional[float]] = []
    for v in rev_row[1:]:
        values.append(_to_float(v))
    yoy: List[Optional[float]] = []
    for i in range(len(values)):
        if i == 0:
            yoy.append(None)
        else:
            prev = values[i - 1]
            cur = values[i]
            if prev is None or cur is None or prev == 0:
                yoy.append(None)
            else:
                yoy.append((cur - prev) / prev * 100.0)
    return ["同比 / YoY"] + [_fmt_yoy(v) for v in yoy]


def _financial_table(doc, data: Dict[str, Any],
                     theme_colors: Dict[str, RGBColor]) -> None:
    ft = data.get("financial_table")
    if not ft:
        return

    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "财务与估值 / Financials & Valuation",
             size=20, bold=True, color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    headers = list(ft.get("headers") or [])
    rows = list(ft.get("rows") or [])
    highlight_rows = list(ft.get("highlight_rows") or [])
    source = ft.get("source") or "Source: Morgan Stanley Research"
    auto_yoy = bool(ft.get("auto_yoy", True))

    # 自动插入 YoY 行（若未显式提供同比行，则在 revenue 之后补一行）
    yoy_row: Optional[List[Any]] = None
    if auto_yoy and not any(_is_growth_row(r) for r in rows):
        yoy_row = _auto_yoy_from_rows(rows)
        if yoy_row is not None and len(yoy_row) >= 2:
            insert_idx = None
            for i, r in enumerate(rows):
                if _is_revenue_row(r):
                    insert_idx = i + 1
                    break
            if insert_idx is not None:
                rows = (rows[:insert_idx] + [yoy_row]
                        + rows[insert_idx:])
                # 由于插入了一行，原 highlight_rows 索引需向后平移
                highlight_rows = [
                    (i + 1) if i >= insert_idx else i
                    for i in highlight_rows
                ]

    if not headers and not rows:
        return

    # ---- 增强功能：如果 data 含 scenarios["base"]，在历史数据右侧增加预测列 ----
    scenarios = data.get("scenarios")
    base_scenario = None
    forecast_years = []
    if scenarios and isinstance(scenarios, dict):
        base_scenario = scenarios.get("base")
    if base_scenario and isinstance(base_scenario, dict):
        forecast_years = list(base_scenario.get("revenue_forecast") or [])

    # 如果有预测数据，扩展 headers 和 rows
    if forecast_years:
        # 推导预测年份标签（基于 headers 中最后一个历史年份）
        last_hist_year = None
        for h in reversed(headers):
            m = re.search(r"FY(\d{2,4})", str(h))
            if m:
                last_hist_year = int(m.group(1))
                break
        if last_hist_year is None:
            last_hist_year = 2025  # 默认
        # 生成预测年份标签
        forecast_labels = [f"FY{last_hist_year + i + 1}E"
                           for i in range(len(forecast_years))]
        # 扩展 headers
        headers = list(headers) + forecast_labels
        # 在 rows 中增加预测列（空值占位，后续按行类型填充）
        for r_idx, row in enumerate(rows):
            if _is_revenue_row(row):
                # 收入行：填入 base 情景的收入预测
                for fv in forecast_years:
                    row.append(f"{fv:,.0f}" if fv else "—")
            else:
                for _ in forecast_years:
                    row.append("")

    n_cols = len(headers) if headers else max(
        (len(r) for r in rows), default=1)
    n_rows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    # 列宽：第一列稍宽，其余均等
    total_width_in = 6.5
    first_w_in = 2.2
    rest_w_in = (total_width_in - first_w_in) / max(1, n_cols - 1)
    widths = [Inches(first_w_in)] + [Inches(rest_w_in)] * (n_cols - 1)

    # 表头
    current_row = 0
    if headers:
        for c_idx, h in enumerate(headers):
            tc = table.rows[current_row].cells[c_idx]
            tc.width = widths[c_idx]
            _set_cell_shd(tc, _hex(theme_colors["navy"]))
            _set_cell_borders(tc, color=theme_colors["navy"], size="8",
                              sides=("top", "bottom"))
            _add_run(tc.paragraphs[0], str(h), size=10, bold=True,
                     color=WHITE)
            tc.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT
                                          if c_idx == 0
                                          else WD_ALIGN_PARAGRAPH.RIGHT)
        current_row += 1

    # 数据行
    # 预测列起始索引（原始 headers 长度，不含扩展的预测列）
    original_header_count = len(headers) - len(forecast_labels) if forecast_years else len(headers)
    for r_idx, row in enumerate(rows):
        is_highlight = r_idx in highlight_rows
        is_margin = _is_margin_row(row)
        is_capex = _is_capex_row(row)
        is_growth = _is_growth_row(row)
        for c_idx in range(n_cols):
            tc = table.rows[current_row + r_idx].cells[c_idx]
            tc.width = widths[c_idx]
            raw = str(row[c_idx]) if c_idx < len(row) else ""

            # 预测列：浅蓝色背景标识
            is_forecast_col = forecast_years and c_idx >= original_header_count

            if r_idx % 2 == 1 and not is_highlight:
                _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))

            # 预测列浅蓝底色
            if is_forecast_col:
                _set_cell_shd(tc, "EBF5FB")

            # highlight_rows：深蓝粗体字，保持纸张干净（不填金色底）
            if is_highlight:
                _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))

            _set_cell_borders(tc, color=theme_colors["line"], size="2",
                              sides=("top", "bottom", "left", "right"))

            para = tc.paragraphs[0]
            para.alignment = (WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0
                              else WD_ALIGN_PARAGRAPH.RIGHT)

            if c_idx == 0:
                # 行首标签
                color = (theme_colors["navy"] if is_highlight
                         else theme_colors["text"])
                _add_run(para, raw, size=10, bold=is_highlight,
                           color=color)
                continue

            cell_text = raw
            # CAPEX：括号化负数
            if is_capex:
                val = _to_float(raw)
                if val is not None:
                    cell_text = f"({abs(val):,.0f})"
            # YoY/同比 行：符号化与红绿强调
            if is_growth and raw not in ("", "—"):
                val = _to_float(raw)
                if val is not None:
                    cell_text = f"{val:+.1f}%"
                    g_color = (RGBColor(0x1F, 0x7A, 0x3E)
                               if val >= 0 else theme_colors["uw"])
                    _add_run(para, cell_text, size=10, bold=True,
                             color=g_color)
                    continue

            # Margin 行：金色强调
            if is_margin:
                _add_run(para, cell_text, size=10, bold=True,
                         color=theme_colors["gold"])
                continue

            # highlight_rows：深蓝粗体
            if is_highlight:
                _add_run(para, cell_text, size=10, bold=True,
                         color=theme_colors["navy"])
                continue

            # 预测列：蓝色字体（输入假设）
            if is_forecast_col and raw not in ("", "—"):
                _add_run(para, cell_text, size=10, bold=True,
                         color=RGBColor(0x1A, 0x5C, 0xB0))
                continue

            _add_run(para, cell_text, size=10,
                     color=theme_colors["text"])

    # 顶部粗线 + 底部粗线
    if headers:
        for c_idx in range(n_cols):
            tc = table.rows[0].cells[c_idx]
            _set_cell_borders(tc, color=theme_colors["navy"], size="16",
                              sides=("top",))
    if rows:
        last_row_idx = current_row + len(rows) - 1
        for c_idx in range(n_cols):
            tc = table.rows[last_row_idx].cells[c_idx]
            _set_cell_borders(tc, color=theme_colors["navy"], size="16",
                              sides=("bottom",))

    # Source 行
    src = doc.add_paragraph()
    _add_run(src, str(source), size=9, italic=True,
             color=theme_colors["muted"])

    doc.add_page_break()


# =============================================================================
# 11. Rating Table 评级总表
# =============================================================================

_RATING_COLOR = {
    "OW": OW_GREEN, "OVERWEIGHT": OW_GREEN, "BUY": OW_GREEN, "买入": OW_GREEN,
    "EW": EW_ORANGE, "EQUALWEIGHT": EW_ORANGE, "HOLD": EW_ORANGE, "持有": EW_ORANGE,
    "UW": UW_RED, "UNDERWEIGHT": UW_RED, "SELL": UW_RED, "卖出": UW_RED,
}


def _rating_table(doc, data: Dict[str, Any],
                  theme_colors: Dict[str, RGBColor]) -> None:
    rt = data.get("rating_table") or []
    if not rt:
        return

    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "评级总表 / Rating Summary", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    # 支持两种格式：
    # 1) list of dicts: [{"metric": "...", "score": "...", "comment": "..."}, ...]
    # 2) list of lists: [["Metric", "Score", "Comment"], [row1], [row2], ...]
    if isinstance(rt[0], dict):
        # dict 格式：自动提取表头和行
        headers_row = ["Metric", "Score", "Comment"]
        data_rows = []
        for item in rt:
            if isinstance(item, dict):
                data_rows.append([
                    str(item.get("metric") or ""),
                    str(item.get("score") or ""),
                    str(item.get("comment") or ""),
                ])
            else:
                data_rows.append([str(item), "", ""])
    else:
        # list 格式：第一行是表头
        headers_row = rt[0]
        data_rows = rt[1:]

    n_cols = len(headers_row)

    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 表头
    for c_idx, h in enumerate(headers_row):
        tc = table.rows[0].cells[c_idx]
        _set_cell_shd(tc, _hex(theme_colors["navy"]))
        _add_run(tc.paragraphs[0], str(h), size=10, bold=True, color=WHITE)

    rating_col = 1  # 默认第二列是 Rating
    for r_idx, row in enumerate(data_rows, start=1):
        for c_idx, val in enumerate(row):
            tc = table.rows[r_idx].cells[c_idx]
            if r_idx % 2 == 0:
                _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
            text = str(val)
            color = theme_colors["text"]
            if c_idx == rating_col:
                color = _RATING_COLOR.get(str(text).upper().strip(),
                                          theme_colors["text"])
                _add_run(tc.paragraphs[0], text, size=10, bold=True,
                         color=color)
            else:
                _add_run(tc.paragraphs[0], text, size=10, color=color)

    doc.add_page_break()


# =============================================================================
# 12. 60 铲子股行业列表
# =============================================================================

def _shovel_stocks(doc, data: Dict[str, Any],
                   theme_colors: Dict[str, RGBColor]) -> None:
    ss = data.get("shovel_stocks") or []
    if not ss:
        return

    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "60 铲子股行业列表 / Shovel Stocks Universe",
             size=20, bold=True, color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    headers = ["#", "公司 / Company", "Ticker", "核心产品 / Product",
               "分析师 / Analyst", "市值 ($mn)", "1年回报 (%)"]
    n_cols = len(headers)

    table = doc.add_table(rows=1 + len(ss), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(0.4), Inches(1.6), Inches(0.9), Inches(1.3),
              Inches(1.0), Inches(0.8), Inches(0.8)]
    for c_idx, w in enumerate(widths):
        table.rows[0].cells[c_idx].width = w

    for c_idx, h in enumerate(headers):
        tc = table.rows[0].cells[c_idx]
        _set_cell_shd(tc, _hex(theme_colors["navy"]))
        _add_run(tc.paragraphs[0], h, size=10, bold=True, color=WHITE)

    for r_idx, row in enumerate(ss):
        row_idx = r_idx + 1
        rank = row.get("rank", r_idx + 1)
        company = str(row.get("company") or "")
        ticker = str(row.get("ticker") or "")
        product = str(row.get("product") or "")
        analyst = str(row.get("analyst") or "")
        try:
            mc = f"{float(row.get('market_cap_mn') or 0):,.1f}"
        except (TypeError, ValueError):
            mc = str(row.get("market_cap_mn") or "")
        try:
            perf = float(row.get("perf_1y_pct") or 0)
            perf_str = f"{perf:+.2f}%"
        except (TypeError, ValueError):
            perf_str = str(row.get("perf_1y_pct") or "")
            perf = 0.0

        values = [rank, company, ticker, product, analyst, mc, perf_str]

        for c_idx, val in enumerate(values):
            tc = table.rows[row_idx].cells[c_idx]
            tc.width = widths[c_idx]
            if r_idx % 2 == 1:
                _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))

            para = tc.paragraphs[0]
            if c_idx in (5, 6):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 回报 > 100% 橙色高亮
            if c_idx == 6 and perf >= 100:
                _set_cell_shd(tc, _hex(theme_colors["ew"]))
                _add_run(para, str(val), size=10, bold=True, color=WHITE)
            elif c_idx == 6 and perf < 0:
                _add_run(para, str(val), size=10, bold=True,
                         color=theme_colors["uw"])
            elif c_idx == 6 and perf >= 0:
                _add_run(para, str(val), size=10, bold=True,
                         color=theme_colors["ow"])
            else:
                _add_run(para, str(val), size=10,
                         color=theme_colors["text"])

    src = doc.add_paragraph()
    _add_run(src, "Source: Morgan Stanley Research  |  注：回报 >=100% 橙色高亮显示",
             size=9, italic=True, color=theme_colors["muted"])

    doc.add_page_break()


# =============================================================================
# 12a. Our Thesis in N Charts（核心论点图表模块）
# =============================================================================

def _thesis_charts_section(doc, data: Dict[str, Any],
                           theme_colors: Dict[str, RGBColor],
                           chart_temp_paths: List[str]) -> None:
    """Our Thesis in N Charts -- 在目录后、正文前，用N个关键图表概括核心论点。

    data["thesis_charts"] = [
        {"title": "图表标题", "chart": {...chart_data...}},
        {"title": "表格标题", "table": {"headers": [...], "rows": [...]}},
    ]
    """
    thesis_charts = data.get("thesis_charts") or []
    if not thesis_charts:
        return

    n = len(thesis_charts)
    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, f"Our Thesis in {n} Charts",
             size=20, bold=True, color=theme_colors["brand_blue"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    for item in thesis_charts:
        title = str(item.get("title") or "Chart")
        exhibit_label = _next_exhibit_label(title)

        # 图表类型
        chart_data = item.get("chart")
        if chart_data and isinstance(chart_data, dict):
            ch_title = doc.add_paragraph()
            _add_run(ch_title, exhibit_label, size=13, bold=True,
                     color=theme_colors["brand_blue"])
            doc.add_paragraph()
            path = _add_chart_image(doc, chart_data, theme_colors)
            if path:
                chart_temp_paths.append(path)
            src = doc.add_paragraph()
            _add_run(src, _exhibit_source(item.get("source")),
                     size=9, italic=True,
                     color=theme_colors["aux_gray"])
            doc.add_paragraph()
            continue

        # 表格类型
        table_data = item.get("table")
        if table_data and isinstance(table_data, dict):
            headers = table_data.get("headers") or []
            rows = table_data.get("rows") or []
            if headers and rows:
                tbl_title = doc.add_paragraph()
                _add_run(tbl_title, exhibit_label, size=13, bold=True,
                         color=theme_colors["brand_blue"])
                doc.add_paragraph()
                _simple_table(doc, [headers] + rows, theme_colors)
                src = doc.add_paragraph()
                _add_run(src, _exhibit_source(item.get("source")),
                         size=9, italic=True,
                         color=theme_colors["aux_gray"])
                doc.add_paragraph()
                continue

        # Metric KPI 卡片类型（type="metric"）
        metric_value = item.get("value")
        metric_subtitle = item.get("subtitle", "")
        if metric_value:
            m_title = doc.add_paragraph()
            _add_run(m_title, exhibit_label, size=13, bold=True,
                     color=theme_colors["brand_blue"])
            # 大号数值
            val_p = doc.add_paragraph()
            _add_run(val_p, str(metric_value), size=28, bold=True,
                     color=theme_colors["brand_blue"])
            if metric_subtitle:
                _add_run(val_p, f"  {metric_subtitle}", size=11,
                         color=theme_colors["aux_gray"])
            doc.add_paragraph()
            continue

        # 纯文本/描述类型
        desc = item.get("description") or item.get("text") or ""
        if desc:
            d_title = doc.add_paragraph()
            _add_run(d_title, exhibit_label, size=13, bold=True,
                     color=theme_colors["brand_blue"])
            _body_paragraph(doc, str(desc), theme_colors)
            doc.add_paragraph()

    doc.add_page_break()


# =============================================================================
# 12b. 产业链/价值链分析
# =============================================================================

def _value_chain_section(doc, data: Dict[str, Any],
                         theme_colors: Dict[str, RGBColor]) -> None:
    """产业链/价值链分析章节。

    data["value_chain"] = [
        {"category": "大脑 (Brain)", "color": "#00559F",
         "companies": [
             {"name": "百度", "ticker": "BIDU", "note": "AI大模型"},
             {"name": "科大讯飞", "ticker": "002230", "note": "语音AI"},
         ]},
        {"category": "身体组件 (Body)", "color": "#3B81B9", ...},
    ]
    """
    value_chain = data.get("value_chain") or []
    if not value_chain:
        return

    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "产业链分析 / Value Chain Analysis",
             size=20, bold=True, color=theme_colors["brand_blue"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    brand_blue = theme_colors.get("brand_blue", MS_BRAND_BLUE)
    text_color = theme_colors["text"]
    muted = theme_colors["muted"]

    for cat_data in value_chain:
        category = str(cat_data.get("category") or "Category")
        cat_color_hex = str(cat_data.get("color") or _hex(brand_blue))
        companies = cat_data.get("companies") or cat_data.get("items") or cat_data.get("stocks") or []

        # 类别标题行（带彩色背景条）
        cat_table = doc.add_table(rows=1, cols=1)
        cat_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cat_cell = cat_table.rows[0].cells[0]
        cat_cell.width = Inches(6.5)
        _set_cell_shd(cat_cell, cat_color_hex.upper())
        _set_cell_borders(cat_cell, color=RGBColor(0xFF, 0xFF, 0xFF),
                          size="0")
        cat_para = cat_cell.paragraphs[0]
        cat_para.paragraph_format.space_before = Pt(4)
        cat_para.paragraph_format.space_after = Pt(4)
        _add_run(cat_para, f"  {category}", size=12, bold=True,
                 color=WHITE)

        # 公司列表 — 使用表格渲染，每公司一行
        if companies:
            comp_table = doc.add_table(rows=1 + len(companies), cols=3)
            comp_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            comp_table.autofit = False
            comp_widths = [Inches(2.2), Inches(0.8), Inches(3.5)]
            for c_idx, w in enumerate(comp_widths):
                comp_table.rows[0].cells[c_idx].width = w

            # 表头
            comp_headers = ["公司 / Company (Ticker)", "评级 / Rating", "描述 / Description"]
            for c_idx, h in enumerate(comp_headers):
                tc = comp_table.rows[0].cells[c_idx]
                _set_cell_shd(tc, _hex(brand_blue))
                _add_run(tc.paragraphs[0], h, size=9, bold=True,
                         color=WHITE, en_title=True)

            for comp_idx, comp in enumerate(companies):
                comp_name = str(comp.get("name") or "")
                ticker = str(comp.get("ticker") or "")
                note = str(comp.get("note") or comp.get("description") or "")
                rating = comp.get("rating", "")

                row_idx = comp_idx + 1
                if comp_idx % 2 == 1:
                    for c in range(3):
                        _set_cell_shd(comp_table.rows[row_idx].cells[c],
                                      _hex(theme_colors["alt_bg"]))

                # 公司名称 + Ticker 列
                tc0 = comp_table.rows[row_idx].cells[0]
                display = comp_name
                if ticker:
                    display += f" ({ticker})"
                _add_run(tc0.paragraphs[0], display, size=10, bold=True,
                         color=text_color)

                # 评级列
                tc1 = comp_table.rows[row_idx].cells[1]
                if rating:
                    r_upper = str(rating).strip().upper()
                    if r_upper in ("O", "OVERWEIGHT"):
                        _add_run(tc1.paragraphs[0], r_upper, size=9, bold=True,
                                 color=theme_colors.get("green_pos", MS_GREEN_POS))
                    elif r_upper in ("U", "UNDERWEIGHT"):
                        _add_run(tc1.paragraphs[0], r_upper, size=9, bold=True,
                                 color=theme_colors.get("red_neg", MS_RED_NEG))
                    elif r_upper in ("E", "EQUAL-WEIGHT", "EQUAL WEIGHT"):
                        _add_run(tc1.paragraphs[0], r_upper, size=9, bold=True,
                                 color=theme_colors.get("ew", EW_ORANGE))
                    else:
                        _add_run(tc1.paragraphs[0], r_upper, size=9,
                                 bold=True, color=muted)
                else:
                    _add_run(tc1.paragraphs[0], "—", size=9, color=muted)

                # 描述列
                tc2 = comp_table.rows[row_idx].cells[2]
                _add_run(tc2.paragraphs[0], note, size=10, color=muted)

        doc.add_paragraph()  # 类别间空行

    src = doc.add_paragraph()
    _add_run(src, "Source: Morgan Stanley Research",
             size=9, italic=True, color=theme_colors["aux_gray"])

    doc.add_page_break()


# =============================================================================
# 12c. 铲子股/概念股列表（MS 标准格式）
# =============================================================================

def _shovel_stocks_list(doc, data: Dict[str, Any],
                        theme_colors: Dict[str, RGBColor]) -> None:
    """铲子股/概念股列表（MS 太空报告风格）。

    data["shovel_stocks_list"] = [
        {"name": "Rocket Lab", "ticker": "RKLB", "rating": "O",
         "price": 66.74, "date": "2026-04-09", "analyst": "Adam Jonas",
         "note": "Small launch vehicle leader"},
    ]
    评级: O=Overweight, E=Equal-weight, U=Underweight, NR=Not-Rated
    """
    stocks = data.get("shovel_stocks_list") or []
    if not stocks:
        return

    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "概念股列表 / Shovel Stocks List",
             size=20, bold=True, color=theme_colors["brand_blue"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    brand_blue = theme_colors.get("brand_blue", MS_BRAND_BLUE)
    text_color = theme_colors["text"]
    muted = theme_colors["muted"]

    headers = ["Company (Ticker)", "Rating (As Of)", "Price (Date)",
               "Analyst", "Note"]
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(stocks), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(1.6), Inches(1.0), Inches(1.0), Inches(1.0),
             Inches(2.2)]
    for c_idx, w in enumerate(widths):
        table.rows[0].cells[c_idx].width = w

    for c_idx, h in enumerate(headers):
        tc = table.rows[0].cells[c_idx]
        _set_cell_shd(tc, _hex(brand_blue))
        _add_run(tc.paragraphs[0], h, size=9, bold=True,
                 color=WHITE, en_title=True)

    for r_idx, stock in enumerate(stocks):
        row_idx = r_idx + 1
        name = str(stock.get("name") or "")
        ticker = str(stock.get("ticker") or "")
        rating = str(stock.get("rating") or "NR").strip().upper()
        price = stock.get("price")
        price_date = str(stock.get("date") or "")
        analyst = str(stock.get("analyst") or "")
        note = str(stock.get("note") or "")

        # 格式化价格
        if price is not None:
            try:
                price_str = _fmt_currency(float(price),
                                         stock.get("currency", "USD"))
            except (TypeError, ValueError):
                price_str = str(price)
        else:
            price_str = ""
        price_display = f"{price_str} ({price_date})" if price_date else price_str

        values = [
            f"{name} ({ticker})" if ticker else name,
            rating,
            price_display,
            analyst,
            note,
        ]

        for c_idx, val in enumerate(values):
            tc = table.rows[row_idx].cells[c_idx]
            tc.width = widths[c_idx]
            if r_idx % 2 == 1:
                _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))

            para = tc.paragraphs[0]

            # 评级列颜色
            if c_idx == 1:
                if rating in ("O", "OVERWEIGHT"):
                    _add_run(para, val, size=9, bold=True,
                             color=theme_colors.get("green_pos",
                                                     MS_GREEN_POS))
                elif rating in ("U", "UNDERWEIGHT"):
                    _add_run(para, val, size=9, bold=True,
                             color=theme_colors.get("red_neg",
                                                     MS_RED_NEG))
                elif rating in ("E", "EQUAL-WEIGHT", "EQUAL WEIGHT"):
                    _add_run(para, val, size=9, bold=True,
                             color=theme_colors.get("ew", EW_ORANGE))
                else:
                    _add_run(para, val, size=9, color=muted)
            else:
                _add_run(para, val, size=9, color=text_color)

    src = doc.add_paragraph()
    _add_run(src,
             "Source: Morgan Stanley Research  |  "
             "Rating: O=Overweight, E=Equal-weight, U=Underweight, "
             "NR=Not-Rated",
             size=9, italic=True, color=theme_colors["aux_gray"])

    doc.add_page_break()


# =============================================================================
# 13. 图表插入页（Exhibit）
# =============================================================================

def _exhibit_page(doc, img_path: str, title: str,
                  theme_colors: Dict[str, RGBColor],
                  source: Optional[str] = None,
                  use_counter: bool = True) -> None:
    """插入 Exhibit 图表页。当 use_counter=True 时自动编号。"""
    display_title = _next_exhibit_label(title) if use_counter else str(title)
    head = doc.add_paragraph()
    _add_run(head, display_title, size=14, bold=True,
             color=theme_colors["brand_blue"])
    doc.add_paragraph()

    try:
        doc.add_picture(img_path, width=Inches(6.5))
    except Exception:
        note = doc.add_paragraph()
        _add_run(note, f"[图表无法加载: {img_path}]", size=10,
                 italic=True, color=theme_colors["muted"])

    src = doc.add_paragraph()
    _add_run(src, _exhibit_source(source),
             size=9, italic=True, color=theme_colors["aux_gray"])

    doc.add_page_break()


def _exhibits_section(doc, data: Dict[str, Any],
                      theme_colors: Dict[str, RGBColor]) -> None:
    exhibits = data.get("exhibits") or []
    for ex in exhibits:
        img = ex.get("image") if isinstance(ex, dict) else ex
        if not img:
            continue
        title = (ex.get("title") if isinstance(ex, dict) else None
                 ) or "Chart"
        source = (ex.get("source") if isinstance(ex, dict) else None
                  ) or None
        _exhibit_page(doc, img, title, theme_colors, source=source,
                      use_counter=True)


# =============================================================================
# 14. Appendix 附录
# =============================================================================

def _appendix(doc, data: Dict[str, Any],
              theme_colors: Dict[str, RGBColor]) -> None:
    appendix = data.get("appendix") or data.get("appendix_table")
    if not appendix:
        return

    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "附录 / Appendix", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    if isinstance(appendix, str):
        for para in [p.strip() for p in re.split(r"\n\s*\n", appendix)
                     if p.strip()]:
            _body_paragraph(doc, para, theme_colors)
    elif isinstance(appendix, list) and appendix:
        if isinstance(appendix[0], list):
            _simple_table(doc, appendix, theme_colors)
        else:
            for item in appendix:
                _body_paragraph(doc, str(item), theme_colors)

    doc.add_page_break()


# =============================================================================
# 15. Disclosure Section 免责声明（中英双语满版 8-9pt）
# =============================================================================

_DISCLOSURE_CN = """\
重要披露：本报告由 Morgan Stanley 研究部门制作，仅供机构客户参考，不构成任何投资建议。\
本报告所载信息来源自公开渠道，Morgan Stanley 对信息的准确性、完整性不作任何明示或默示保证。\
报告中的观点仅代表分析师个人判断，可能与 Morgan Stanley 其他业务部门的立场不同。\
投资涉及风险，过去的表现不代表未来结果。投资者应根据自身判断独立评估本报告内容，并承担相应风险。\
未经 Morgan Stanley 书面授权，任何机构或个人不得以任何形式复制、传播、引用本报告的任何内容。\
本报告在不同司法管辖区的分发可能受当地法律法规限制，接收者应遵守当地法律。
"""

_DISCLOSURE_EN = """\
Important Disclosures: This research note is produced by the Morgan Stanley \
research department and is intended solely for institutional clients. It does \
not constitute investment advice. Information contained herein is derived from \
public sources; Morgan Stanley makes no representation or warranty, express or \
implied, as to its accuracy or completeness. Views expressed reflect the \
analyst's current judgment and may differ from the views of other Morgan \
Stanley business areas. All investments involve risk; past performance is not \
indicative of future results. Investors should independently evaluate the \
information herein and rely on their own judgment. No part of this report may \
be reproduced, redistributed or cited in any form without written permission \
from Morgan Stanley. Distribution of this report in certain jurisdictions may \
be restricted by local law. Recipients should comply with applicable local \
regulations.
"""


def _disclosure(doc, data: Dict[str, Any],
                theme_colors: Dict[str, RGBColor]) -> None:
    head = doc.add_paragraph(style='Heading 1')
    _add_run(head, "免责声明 / Disclosures", size=18, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    text_cn = data.get("disclosure_cn") or _DISCLOSURE_CN
    text_en = data.get("disclosure_en") or _DISCLOSURE_EN

    cn_head = doc.add_paragraph()
    _add_run(cn_head, "中文 / Chinese", size=11, bold=True,
             color=theme_colors["navy"])
    for para in [p.strip() for p in re.split(r"\n\s*\n", text_cn)
                 if p.strip()]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        _add_run(p, para, size=9, color=theme_colors["text"])

    doc.add_paragraph()
    en_head = doc.add_paragraph()
    _add_run(en_head, "English", size=11, bold=True,
             color=theme_colors["navy"])
    for para in [p.strip() for p in re.split(r"\n\s*\n", text_en)
                 if p.strip()]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        _add_run(p, para, size=9, color=theme_colors["text"],
                 en_title=False)


# =============================================================================
# 16. 情景分析 / Scenario Comparison
# =============================================================================

def _scenario_comparison_section(doc, data: Dict[str, Any],
                                theme_colors: Dict[str, RGBColor],
                                language: str = "zh") -> None:
    """三情景对比表格：熊市 / 基准 / 牛市，基准列金色高亮。"""
    scenarios = data.get("scenarios")
    if not scenarios or not isinstance(scenarios, dict):
        return
    bear = scenarios.get("bear")
    base = scenarios.get("base")
    bull = scenarios.get("bull")
    if not (bear and base and bull):
        return

    currency = data.get("currency", "USD")
    sym = {"USD": "$", "CNY": "¥", "HKD": "HK$", "EUR": "€", "GBP": "£"}.get(currency, "$")

    # 标题
    head = doc.add_paragraph(style='Heading 1')
    head.paragraph_format.space_before = Pt(8)
    _add_run(head, "三、情景分析 / Scenario Analysis", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    # 构建表头和数据行
    headers = ["指标 / Metric", "熊市 / Bear", "基准 / Base", "牛市 / Bull"]
    rows_data = [
        ["关键假设 / Key Assumptions", "", "", ""],
        ["WACC", _fmt_percent((bear.get("wacc_value") or 0) * 100),
         _fmt_percent((base.get("wacc_value") or 0) * 100),
         _fmt_percent((bull.get("wacc_value") or 0) * 100)],
        ["Terminal Growth Rate", _fmt_percent((bear.get("terminal_growth_rate") or 0) * 100),
         _fmt_percent((base.get("terminal_growth_rate") or 0) * 100),
         _fmt_percent((bull.get("terminal_growth_rate") or 0) * 100)],
        ["Exit Multiple (EBITDA)", f"{bear.get('exit_multiple_ebitda', '—')}x",
         f"{base.get('exit_multiple_ebitda', '—')}x",
         f"{bull.get('exit_multiple_ebitda', '—')}x"],
        ["Revenue CAGR", _fmt_percent((bear.get("revenue_cagr") or 0) * 100),
         _fmt_percent((base.get("revenue_cagr") or 0) * 100),
         _fmt_percent((bull.get("revenue_cagr") or 0) * 100)],
        ["", "", "", ""],
        ["估值结果 / Valuation", "", "", ""],
        ["Enterprise Value", _fmt_currency(bear.get("enterprise_value"), currency),
         _fmt_currency(base.get("enterprise_value"), currency),
         _fmt_currency(bull.get("enterprise_value"), currency)],
        ["Equity Value", _fmt_currency(bear.get("equity_value"), currency),
         _fmt_currency(base.get("equity_value"), currency),
         _fmt_currency(bull.get("equity_value"), currency)],
        ["Per Share", _fmt_currency(bear.get("per_share_value"), currency),
         _fmt_currency(base.get("per_share_value"), currency),
         _fmt_currency(bull.get("per_share_value"), currency)],
        ["EV/Revenue", f"{bear.get('ev_revenue', '—')}x",
         f"{base.get('ev_revenue', '—')}x",
         f"{bull.get('ev_revenue', '—')}x"],
        ["EV/EBITDA", f"{bear.get('ev_ebitda', '—')}x",
         f"{base.get('ev_ebitda', '—')}x",
         f"{bull.get('ev_ebitda', '—')}x"],
    ]

    n_cols = len(headers)
    n_data_rows = len(rows_data)
    table = doc.add_table(rows=1 + n_data_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total_w = 6.5
    widths = [Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.5)]
    for c_idx, w in enumerate(widths):
        for r_idx in range(1 + n_data_rows):
            table.rows[r_idx].cells[c_idx].width = w

    # 写入表头
    header_colors = [
        _hex(theme_colors["navy"]),
        _hex(theme_colors["bear_color"]),
        _hex(theme_colors["base_color"]),
        _hex(theme_colors["bull_color"]),
    ]
    for c_idx, h in enumerate(headers):
        tc = table.rows[0].cells[c_idx]
        _set_cell_shd(tc, header_colors[c_idx])
        _set_cell_borders(tc, color=theme_colors["navy"], size="8",
                          sides=("top", "bottom"))
        para = tc.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(para, h, size=10, bold=True, color=WHITE)

    # 写入数据行
    section_header_rows = {0, 6}  # "关键假设" 和 "估值结果" 分组标题行
    for r_idx, row in enumerate(rows_data, start=1):
        for c_idx, val in enumerate(row):
            tc = table.rows[r_idx].cells[c_idx]
            para = tc.paragraphs[0]

            # 分组标题行（关键假设 / 估值结果）
            if r_idx - 1 in section_header_rows:
                _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
                _set_cell_borders(tc, color=theme_colors["line"], size="2",
                                  sides=("top", "bottom"))
                if c_idx == 0:
                    _add_run(para, val, size=10, bold=True,
                             color=theme_colors["navy"])
                continue

            # 基准列（c_idx=2）浅绿背景
            if c_idx == 2:
                _set_cell_shd(tc, _hex(theme_colors["base_bg"]))
            elif c_idx == 1:
                # 熊市列浅橙背景（交替行）
                if r_idx % 2 == 0:
                    _set_cell_shd(tc, _hex(theme_colors["bear_bg"]))
            elif c_idx == 3:
                # 牛市列浅紫背景（交替行）
                if r_idx % 2 == 0:
                    _set_cell_shd(tc, _hex(theme_colors["bull_bg"]))
            elif r_idx % 2 == 0:
                _set_cell_shd(tc, _hex(theme_colors["alt_row"]))

            _set_cell_borders(tc, color=theme_colors["line"], size="2",
                              sides=("top", "bottom", "left", "right"))

            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            if c_idx == 0:
                _add_run(para, val, size=10, color=theme_colors["text"])
            elif c_idx == 2:
                # 基准列数值用深绿加粗
                _add_run(para, val, size=10, bold=True,
                         color=theme_colors["base_color"])
            else:
                _add_run(para, val, size=10, color=theme_colors["text"])

    # 顶部和底部粗线
    for c_idx in range(n_cols):
        _set_cell_borders(table.rows[0].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("top",))
        _set_cell_borders(table.rows[n_data_rows].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("bottom",))

    # 分析文字
    doc.add_paragraph()
    analysis = data.get("scenario_analysis_text")
    if analysis:
        _body_paragraph(doc, analysis, theme_colors)
    else:
        # 默认分析文字
        base_ps = base.get("per_share_value", "—")
        bear_ps = bear.get("per_share_value", "—")
        bull_ps = bull.get("per_share_value", "—")
        default_text = (
            f"基准情景下每股价值为 {sym}{base_ps}，"
            f"熊市情景为 {sym}{bear_ps}（下行风险），"
            f"牛市情景为 {sym}{bull_ps}（上行空间）。"
            "基准情景假设反映了当前市场共识与公司基本面趋势，"
            "WACC 与终端增长率取行业中位数水平。"
        )
        _body_paragraph(doc, default_text, theme_colors)

    doc.add_page_break()


# =============================================================================
# 17. WACC 分析 / WACC Analysis
# =============================================================================

def _wacc_section(doc, data: Dict[str, Any],
                  theme_colors: Dict[str, RGBColor],
                  language: str = "zh") -> None:
    """WACC 参数表格 + matplotlib 堆叠条形图。"""
    scenarios = data.get("scenarios")
    if not scenarios or not isinstance(scenarios, dict):
        return
    bear = scenarios.get("bear")
    base = scenarios.get("base")
    bull = scenarios.get("bull")
    if not (bear and base and bull):
        return

    # 标题
    head = doc.add_paragraph(style='Heading 1')
    head.paragraph_format.space_before = Pt(8)
    _add_run(head, "四、WACC 分析 / WACC Analysis", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    # WACC 参数表格
    wacc_headers = ["参数 / Parameter", "熊市 / Bear", "基准 / Base", "牛市 / Bull"]
    wacc_rows = [
        ["Risk-Free Rate (Rf)", _fmt_percent((bear["wacc"].get("rf") or 0) * 100),
         _fmt_percent((base["wacc"].get("rf") or 0) * 100),
         _fmt_percent((bull["wacc"].get("rf") or 0) * 100)],
        ["Equity Risk Premium (ERP)", _fmt_percent((bear["wacc"].get("erp") or 0) * 100),
         _fmt_percent((base["wacc"].get("erp") or 0) * 100),
         _fmt_percent((bull["wacc"].get("erp") or 0) * 100)],
        ["Beta (Levered)", _fmt_number(bear["wacc"].get("beta")),
         _fmt_number(base["wacc"].get("beta")),
         _fmt_number(bull["wacc"].get("beta"))],
        ["Size Premium", _fmt_percent((bear["wacc"].get("size_premium") or 0) * 100),
         _fmt_percent((base["wacc"].get("size_premium") or 0) * 100),
         _fmt_percent((bull["wacc"].get("size_premium") or 0) * 100)],
        ["Country Risk Premium", _fmt_percent((bear["wacc"].get("country_risk") or 0) * 100),
         _fmt_percent((base["wacc"].get("country_risk") or 0) * 100),
         _fmt_percent((bull["wacc"].get("country_risk") or 0) * 100)],
        ["Cost of Equity (Ke)", _fmt_percent((bear["wacc"].get("ke") or 0) * 100),
         _fmt_percent((base["wacc"].get("ke") or 0) * 100),
         _fmt_percent((bull["wacc"].get("ke") or 0) * 100)],
        ["Cost of Debt (Kd)", _fmt_percent((bear["wacc"].get("kd") or 0) * 100),
         _fmt_percent((base["wacc"].get("kd") or 0) * 100),
         _fmt_percent((bull["wacc"].get("kd") or 0) * 100)],
        ["Tax Rate", _fmt_percent((bear["wacc"].get("tax_rate") or 0) * 100),
         _fmt_percent((base["wacc"].get("tax_rate") or 0) * 100),
         _fmt_percent((bull["wacc"].get("tax_rate") or 0) * 100)],
        ["Equity Weight", _fmt_percent((bear["wacc"].get("e_weight") or 0) * 100),
         _fmt_percent((base["wacc"].get("e_weight") or 0) * 100),
         _fmt_percent((bull["wacc"].get("e_weight") or 0) * 100)],
        ["Debt Weight", _fmt_percent((bear["wacc"].get("d_weight") or 0) * 100),
         _fmt_percent((base["wacc"].get("d_weight") or 0) * 100),
         _fmt_percent((bull["wacc"].get("d_weight") or 0) * 100)],
        ["WACC", _fmt_percent((bear.get("wacc_value") or 0) * 100),
         _fmt_percent((base.get("wacc_value") or 0) * 100),
         _fmt_percent((bull.get("wacc_value") or 0) * 100)],
    ]

    n_cols = len(wacc_headers)
    n_rows = len(wacc_rows)
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.5)]
    for c_idx, w in enumerate(widths):
        for r_idx in range(n_rows + 1):
            table.rows[r_idx].cells[c_idx].width = w

    # 表头
    wacc_header_colors = [
        _hex(theme_colors["navy"]),
        _hex(theme_colors["bear_color"]),
        _hex(theme_colors["base_color"]),
        _hex(theme_colors["bull_color"]),
    ]
    for c_idx, h in enumerate(wacc_headers):
        tc = table.rows[0].cells[c_idx]
        _set_cell_shd(tc, wacc_header_colors[c_idx])
        _set_cell_borders(tc, color=theme_colors["navy"], size="8",
                          sides=("top", "bottom"))
        para = tc.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(para, h, size=10, bold=True, color=WHITE)

    # 数据行
    for r_idx, row in enumerate(wacc_rows):
        is_wacc_row = (r_idx == len(wacc_rows) - 1)  # 最后一行是 WACC 汇总
        for c_idx, val in enumerate(row):
            tc = table.rows[r_idx + 1].cells[c_idx]
            para = tc.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT

            if is_wacc_row:
                _set_cell_shd(tc, _hex(theme_colors["summary_bg"]))
                _set_cell_borders(tc, color=theme_colors["navy"], size="4",
                                  sides=("top", "bottom", "left", "right"))
                if c_idx == 0:
                    _add_run(para, val, size=10, bold=True,
                             color=theme_colors["navy"])
                elif c_idx == 2:
                    _add_run(para, val, size=11, bold=True,
                             color=theme_colors["base_color"])
                else:
                    _add_run(para, val, size=10, bold=True,
                             color=theme_colors["text"])
            else:
                if c_idx == 2:
                    _set_cell_shd(tc, _hex(theme_colors["base_bg"]))
                elif c_idx == 1 and r_idx % 2 == 1:
                    _set_cell_shd(tc, _hex(theme_colors["bear_bg"]))
                elif c_idx == 3 and r_idx % 2 == 1:
                    _set_cell_shd(tc, _hex(theme_colors["bull_bg"]))
                elif r_idx % 2 == 1:
                    _set_cell_shd(tc, _hex(theme_colors["alt_row"]))
                _set_cell_borders(tc, color=theme_colors["line"], size="2",
                                  sides=("top", "bottom", "left", "right"))
                if c_idx == 0:
                    _add_run(para, val, size=10, color=theme_colors["text"])
                elif c_idx == 2:
                    _add_run(para, val, size=10, bold=True,
                             color=theme_colors["base_color"])
                else:
                    _add_run(para, val, size=10, color=theme_colors["text"])

    # 顶部和底部粗线
    for c_idx in range(n_cols):
        _set_cell_borders(table.rows[0].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("top",))
        _set_cell_borders(table.rows[n_rows].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("bottom",))

    # WACC 构成堆叠条形图
    doc.add_paragraph()
    plt, ok = _try_import_matplotlib()
    if ok:
        # Apply publication-grade global style
        _setup_publication_style(plt)

        try:
            fig, ax = plt.subplots(figsize=(6.2, 3.2), dpi=300)
            fig.patch.set_facecolor('#FFFFFF')
            ax.set_facecolor('#FFFFFF')

            labels_list = ["Bear", "Base", "Bull"]
            ke_vals = [
                _to_float(bear["wacc"].get("ke")) or 0,
                _to_float(base["wacc"].get("ke")) or 0,
                _to_float(bull["wacc"].get("ke")) or 0,
            ]
            kd_vals = [
                _to_float(bear["wacc"].get("kd")) or 0,
                _to_float(base["wacc"].get("kd")) or 0,
                _to_float(bull["wacc"].get("kd")) or 0,
            ]
            e_weights = [
                _to_float(bear["wacc"].get("e_weight")) or 0.8,
                _to_float(base["wacc"].get("e_weight")) or 0.8,
                _to_float(bull["wacc"].get("e_weight")) or 0.8,
            ]
            d_weights = [
                _to_float(bear["wacc"].get("d_weight")) or 0.2,
                _to_float(base["wacc"].get("d_weight")) or 0.2,
                _to_float(bull["wacc"].get("d_weight")) or 0.2,
            ]

            # Ke contribution = Ke * equity weight; Kd contribution = Kd * debt weight * (1-tax)
            tax_rates = [
                _to_float(bear["wacc"].get("tax_rate")) or 0.25,
                _to_float(base["wacc"].get("tax_rate")) or 0.25,
                _to_float(bull["wacc"].get("tax_rate")) or 0.25,
            ]
            ke_contrib = [ke_vals[i] * e_weights[i] for i in range(3)]
            kd_contrib = [kd_vals[i] * d_weights[i] * (1 - tax_rates[i]) for i in range(3)]

            xs = list(range(3))
            w = 0.5
            # Rounded bar effect via white edge
            ax.bar(xs, ke_contrib, width=w, color=MS_PALETTE[0],
                   label="Ke Contribution", edgecolor='white', linewidth=0, zorder=3)
            ax.bar(xs, kd_contrib, width=w, bottom=ke_contrib,
                   color=MS_PALETTE[1], label="Kd After-Tax Contribution",
                   edgecolor='white', linewidth=0, zorder=3)

            # 在柱顶标注 WACC 总值
            wacc_vals = [
                (_to_float(bear.get("wacc_value")) or 0) * 100,
                (_to_float(base.get("wacc_value")) or 0) * 100,
                (_to_float(bull.get("wacc_value")) or 0) * 100,
            ]
            totals = [ke_contrib[i] + kd_contrib[i] for i in range(3)]
            for i, (x, total) in enumerate(zip(xs, totals)):
                ax.text(x, total + 0.002, f"{wacc_vals[i]:.1f}%",
                        ha="center", va="bottom", fontsize=10, fontweight="bold",
                        color=MS_CHART_TITLE_COLOR, zorder=5)
                # Component breakdown annotation below WACC value
                ax.text(x, total - 0.008,
                        f"Ke={ke_vals[i]*100:.1f}%, Kd={kd_vals[i]*100:.1f}%",
                        ha="center", va="top", fontsize=6.5,
                        color='#4A4A4A', fontstyle='italic', zorder=5)

            ax.set_xticks(xs)
            ax.set_xticklabels(labels_list)
            ax.set_ylabel("WACC Components", fontsize=9, color=MS_CHART_TICK_COLOR)
            ax.legend(frameon=False, fontsize=8, labelcolor=MS_CHART_TICK_COLOR,
                      loc='upper right')
            # Apply publication style + source watermark
            _apply_ax_publication_style(ax, title="WACC Composition by Scenario")
            _add_source_watermark(fig)
            fig.tight_layout(pad=1.5)

            fd, path = tempfile.mkstemp(prefix="ms_wacc_", suffix=".png")
            os.close(fd)
            try:
                fig.savefig(path, dpi=300, facecolor='#FFFFFF',
                            bbox_inches='tight')
                doc.add_picture(path, width=Inches(6.2))
            except Exception:
                pass
            finally:
                plt.close(fig)
                try:
                    os.remove(path)
                except OSError:
                    pass
        except Exception:
            plt.close("all")

    # 分析文字
    doc.add_paragraph()
    wacc_analysis = data.get("wacc_analysis_text")
    if wacc_analysis:
        _body_paragraph(doc, wacc_analysis, theme_colors)
    else:
        base_wacc = (base.get("wacc_value") or 0) * 100
        default_text = (
            f"基准情景 WACC 为 {base_wacc:.1f}%，"
            "基于当前无风险利率水平、行业 ERP 及公司特定 Beta 调整。"
            "熊市情景采用更高的风险溢价与 Beta，反映下行风险加大时的投资者要求回报率上升。"
        )
        _body_paragraph(doc, default_text, theme_colors)

    doc.add_page_break()


# =============================================================================
# 18. 敏感性分析 / Sensitivity Analysis
# =============================================================================

def _sensitivity_section(doc, data: Dict[str, Any],
                         theme_colors: Dict[str, RGBColor],
                         language: str = "zh") -> None:
    """WACC x TGR -> Per Share Value 矩阵表格，基准位置金色标注。"""
    sens = data.get("sensitivity")
    if not sens or not isinstance(sens, dict):
        return

    wacc_range = sens.get("wacc_range") or []
    tgr_range = sens.get("tgr_range") or []
    matrix = sens.get("matrix") or []
    base_wacc_idx = sens.get("base_wacc_idx", 0)
    base_tgr_idx = sens.get("base_tgr_idx", 0)

    if not wacc_range or not tgr_range or not matrix:
        return

    currency = data.get("currency", "USD")
    sym = {"USD": "$", "CNY": "¥", "HKD": "HK$", "EUR": "€", "GBP": "£"}.get(currency, "$")

    # 标题
    head = doc.add_paragraph(style='Heading 1')
    head.paragraph_format.space_before = Pt(8)
    _add_run(head, "五、敏感性分析 / Sensitivity Analysis", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    _add_run(subtitle, "WACC × Terminal Growth Rate → Per Share Value",
             size=12, italic=True, color=theme_colors["muted"])
    doc.add_paragraph()

    # 表格: 左上角空白，第一行为 TGR，第一列为 WACC
    n_tgr = len(tgr_range)
    n_wacc = len(wacc_range)
    n_cols = n_tgr + 1
    n_rows = n_wacc + 1

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    col_w = min(1.0, 6.0 / n_cols)
    first_col_w = 1.2
    widths = [Inches(first_col_w)] + [Inches(col_w)] * n_tgr
    for c_idx, w in enumerate(widths):
        for r_idx in range(n_rows):
            table.rows[r_idx].cells[c_idx].width = w

    # 左上角单元格
    corner = table.rows[0].cells[0]
    _set_cell_shd(corner, _hex(theme_colors["navy"]))
    _add_run(corner.paragraphs[0], "WACC \\ TGR", size=9, bold=True,
             color=WHITE)
    corner.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第一行：TGR 值
    for j, tgr in enumerate(tgr_range):
        tc = table.rows[0].cells[j + 1]
        _set_cell_shd(tc, _hex(theme_colors["navy"]))
        _set_cell_borders(tc, color=theme_colors["navy"], size="8",
                          sides=("top", "bottom"))
        para = tc.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(para, _fmt_percent((_to_float(tgr) or 0) * 100), size=9, bold=True, color=WHITE)

    # 第一列：WACC 值
    for i, wacc in enumerate(wacc_range):
        tc = table.rows[i + 1].cells[0]
        _set_cell_shd(tc, _hex(theme_colors["navy"]))
        _set_cell_borders(tc, color=theme_colors["navy"], size="4",
                          sides=("left", "right"))
        para = tc.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(para, _fmt_percent((_to_float(wacc) or 0) * 100), size=9, bold=True, color=WHITE)

    # 矩阵数据：颜色编码（深色=高值）
    # 先收集所有有效值用于归一化
    all_values = []
    for row in matrix:
        for v in row:
            fv = _to_float(v)
            if fv is not None:
                all_values.append(fv)
    min_val = min(all_values) if all_values else 0
    max_val = max(all_values) if all_values else 1

    for i, row in enumerate(matrix):
        for j, val in enumerate(row):
            if i >= n_wacc or j >= n_tgr:
                continue
            tc = table.rows[i + 1].cells[j + 1]
            para = tc.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            fv = _to_float(val)
            is_base = (i == base_wacc_idx and j == base_tgr_idx)

            if is_base:
                # 基准位置：黄色背景 + 粗边框
                _set_cell_shd(tc, "FFFF00")
                _set_cell_borders(tc, color=theme_colors["deep_navy"], size="12",
                                  sides=("top", "bottom", "left", "right"))
                _add_run(para, f"{sym}{fv:,.1f}" if fv is not None else "—",
                         size=10, bold=True, color=theme_colors["deep_navy"])
            else:
                # 颜色编码：值越高颜色越深（navy 渐变）
                if fv is not None and max_val > min_val:
                    ratio = (fv - min_val) / (max_val - min_val)
                    # 从 bull_bg（低值）到 deep_navy（高值）
                    r_c = int(0xDD - ratio * (0xDD - 0x1F))
                    g_c = int(0xEE - ratio * (0xEE - 0x38))
                    b_c = int(0xFF - ratio * (0xFF - 0x64))
                    bg_hex = f"{r_c:02X}{g_c:02X}{b_c:02X}"
                    _set_cell_shd(tc, bg_hex)
                    # 高值白色字体，低值深色字体
                    text_color = WHITE if ratio > 0.6 else theme_colors["text"]
                else:
                    _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
                    text_color = theme_colors["text"]

                _set_cell_borders(tc, color=theme_colors["line"], size="2",
                                  sides=("top", "bottom", "left", "right"))
                _add_run(para, f"{sym}{fv:,.1f}" if fv is not None else "—",
                         size=10, bold=True, color=text_color)

    # 顶部和底部粗线
    for c_idx in range(n_cols):
        _set_cell_borders(table.rows[0].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("top",))
        _set_cell_borders(table.rows[n_rows - 1].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("bottom",))

    # 分析文字
    doc.add_paragraph()
    sens_analysis = data.get("sensitivity_analysis_text")
    if sens_analysis:
        _body_paragraph(doc, sens_analysis, theme_colors)
    else:
        default_text = (
            "敏感性分析显示，WACC 每变动 100bps 对每股价值影响显著。"
            "终端增长率假设对长期估值有放大效应，"
            "建议投资者关注利率环境变化对公司资本成本的潜在影响。"
        )
        _body_paragraph(doc, default_text, theme_colors)

    doc.add_page_break()


# =============================================================================
# 19. 可比公司分析 / Comparable Companies
# =============================================================================

def _comps_section(doc, data: Dict[str, Any],
                   theme_colors: Dict[str, RGBColor],
                   language: str = "zh") -> None:
    """可比公司表格，含中位数/均值汇总行。"""
    comps = data.get("comparable_companies")
    if not comps or not isinstance(comps, list):
        return

    currency = data.get("currency", "USD")
    sym = {"USD": "$", "CNY": "¥", "HKD": "HK$", "EUR": "€", "GBP": "£"}.get(currency, "$")

    # 标题
    head = doc.add_paragraph(style='Heading 1')
    head.paragraph_format.space_before = Pt(8)
    _add_run(head, "六、可比公司分析 / Comparable Companies", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    comp_headers = [
        "公司 / Company", "代码 / Ticker",
        f"市值 ({sym}B)", f"EV ({sym}B)",
        "EV/Rev LTM", "EV/Rev NTM",
        "EV/EBITDA LTM", "EV/EBITDA NTM",
        "CAGR 3Y", "备注 / Note"
    ]
    n_cols = len(comp_headers)

    # 计算中位数和均值
    def _calc_stats(key: str) -> Tuple[Optional[float], Optional[float]]:
        vals = [_to_float(c.get(key)) for c in comps]
        vals = [v for v in vals if v is not None]
        if not vals:
            return None, None
        median = sorted(vals)[len(vals) // 2]
        mean = sum(vals) / len(vals)
        return median, mean

    stats = {}
    for key in ("mcap", "ev", "ev_rev_ltm", "ev_rev_ntm",
                "ev_ebitda_ltm", "ev_ebitda_ntm", "cagr_3y"):
        stats[key] = _calc_stats(key)

    n_rows = len(comps) + 2  # comps + median + mean
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(1.1), Inches(0.7), Inches(0.55), Inches(0.55),
              Inches(0.55), Inches(0.55), Inches(0.6), Inches(0.6),
              Inches(0.45), Inches(0.8)]

    for c_idx, w in enumerate(widths):
        for r_idx in range(n_rows + 1):
            table.rows[r_idx].cells[c_idx].width = w

    # 表头
    for c_idx, h in enumerate(comp_headers):
        tc = table.rows[0].cells[c_idx]
        _set_cell_shd(tc, _hex(theme_colors["navy"]))
        _set_cell_borders(tc, color=theme_colors["navy"], size="8",
                          sides=("top", "bottom"))
        para = tc.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(para, h, size=8, bold=True, color=WHITE)

    # 数据行
    for r_idx, comp in enumerate(comps):
        row_data = [
            comp.get("name", ""),
            comp.get("ticker", ""),
            _fmt_number(comp.get("mcap")),
            _fmt_number(comp.get("ev")),
            _fmt_number(comp.get("ev_rev_ltm")),
            _fmt_number(comp.get("ev_rev_ntm")),
            _fmt_number(comp.get("ev_ebitda_ltm")),
            _fmt_number(comp.get("ev_ebitda_ntm")),
            _fmt_percent(comp.get("cagr_3y")),
            comp.get("note", ""),
        ]
        for c_idx, val in enumerate(row_data):
            tc = table.rows[r_idx + 1].cells[c_idx]
            para = tc.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 9) else WD_ALIGN_PARAGRAPH.RIGHT
            if r_idx % 2 == 1:
                _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
            _set_cell_borders(tc, color=theme_colors["line"], size="2",
                              sides=("top", "bottom", "left", "right"))
            _add_run(para, str(val), size=9, color=theme_colors["text"])

    # 中位数行
    median_row_idx = len(comps) + 1
    tc = table.rows[median_row_idx].cells[0]
    _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
    _set_cell_borders(tc, color=theme_colors["navy"], size="4",
                      sides=("top", "bottom"))
    _add_run(tc.paragraphs[0], "中位数 / Median", size=9, bold=True,
             color=theme_colors["navy"])
    for c_idx in range(1, n_cols):
        tc = table.rows[median_row_idx].cells[c_idx]
        _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
        _set_cell_borders(tc, color=theme_colors["navy"], size="4",
                          sides=("top", "bottom"))
        para = tc.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        key_map = {1: "ticker", 2: "mcap", 3: "ev", 4: "ev_rev_ltm",
                   5: "ev_rev_ntm", 6: "ev_ebitda_ltm", 7: "ev_ebitda_ntm",
                   8: "cagr_3y", 9: None}
        key = key_map.get(c_idx)
        if key and stats.get(key):
            med_val = stats[key][0]
            if key in ("cagr_3y",):
                _add_run(para, _fmt_percent(med_val), size=9, bold=True,
                         color=theme_colors["navy"])
            else:
                _add_run(para, _fmt_number(med_val), size=9, bold=True,
                         color=theme_colors["navy"])
        else:
            _add_run(para, "", size=9)

    # 均值行
    mean_row_idx = len(comps) + 2
    tc = table.rows[mean_row_idx].cells[0]
    _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
    _set_cell_borders(tc, color=theme_colors["navy"], size="4",
                      sides=("top", "bottom"))
    _add_run(tc.paragraphs[0], "均值 / Mean", size=9, bold=True,
             color=theme_colors["navy"])
    for c_idx in range(1, n_cols):
        tc = table.rows[mean_row_idx].cells[c_idx]
        _set_cell_shd(tc, _hex(theme_colors["alt_bg"]))
        _set_cell_borders(tc, color=theme_colors["navy"], size="4",
                          sides=("top", "bottom"))
        para = tc.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        key_map = {1: "ticker", 2: "mcap", 3: "ev", 4: "ev_rev_ltm",
                   5: "ev_rev_ntm", 6: "ev_ebitda_ltm", 7: "ev_ebitda_ntm",
                   8: "cagr_3y", 9: None}
        key = key_map.get(c_idx)
        if key and stats.get(key):
            mean_val = stats[key][1]
            if key in ("cagr_3y",):
                _add_run(para, _fmt_percent(mean_val), size=9, bold=True,
                         color=theme_colors["navy"])
            else:
                _add_run(para, _fmt_number(mean_val), size=9, bold=True,
                         color=theme_colors["navy"])
        else:
            _add_run(para, "", size=9)

    # 顶部和底部粗线
    for c_idx in range(n_cols):
        _set_cell_borders(table.rows[0].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("top",))
        _set_cell_borders(table.rows[n_rows].cells[c_idx],
                          color=theme_colors["navy"], size="16",
                          sides=("bottom",))

    # 分析文字
    doc.add_paragraph()
    comps_analysis = data.get("comps_analysis_text")
    if comps_analysis:
        _body_paragraph(doc, comps_analysis, theme_colors)
    else:
        default_text = (
            "可比公司分析显示，目标公司相对同业在 EV/Revenue 与 EV/EBITDA "
            "倍数上处于合理区间。中位数倍数可作为交叉验证 DCF 估值的参考基准。"
        )
        _body_paragraph(doc, default_text, theme_colors)

    doc.add_page_break()


# =============================================================================
# 20. 估值桥 / Valuation Bridge
# =============================================================================

def _valuation_bridge_section(doc, data: Dict[str, Any],
                              theme_colors: Dict[str, RGBColor],
                              language: str = "zh") -> None:
    """瀑布图：PV of FCF → +GGM TV → +Exit Multiple TV → -Net Debt → = Equity Value。"""
    scenarios = data.get("scenarios")
    if not scenarios or not isinstance(scenarios, dict):
        return
    bear = scenarios.get("bear")
    base = scenarios.get("base")
    bull = scenarios.get("bull")
    if not (bear and base and bull):
        return

    currency = data.get("currency", "USD")
    sym = {"USD": "$", "CNY": "¥", "HKD": "HK$", "EUR": "€", "GBP": "£"}.get(currency, "$")

    # 标题
    head = doc.add_paragraph(style='Heading 1')
    head.paragraph_format.space_before = Pt(8)
    _add_run(head, "七、估值桥 / Valuation Bridge", size=20, bold=True,
             color=theme_colors["navy"])
    _gold_short_underline(doc, theme_colors)
    doc.add_paragraph()

    # 使用 matplotlib 生成瀑布图（三情景并排）
    plt, ok = _try_import_matplotlib()
    if ok:
        # Apply publication-grade global style
        _setup_publication_style(plt)

        try:
            fig, ax = plt.subplots(figsize=(6.5, 3.8), dpi=300)
            fig.patch.set_facecolor('#FFFFFF')
            ax.set_facecolor('#FFFFFF')

            labels = ["PV of\nFCF", "+GGM\nTV", "+Exit Mult.\nTV", "-Net\nDebt", "Equity\nValue"]
            scenario_list = [("Bear", bear), ("Base", base), ("Bull", bull)]
            n_bars = len(labels)
            n_groups = len(scenario_list)
            group_width = 0.8
            bar_width = group_width / n_groups

            colors_scenario = [MS_PALETTE[5], MS_PALETTE[0], MS_PALETTE[1]]  # Red, Navy, Blue

            for g_idx, (s_name, s_data) in enumerate(scenario_list):
                pv_fcf = _to_float(s_data.get("pv_fcf")) or 0
                pv_ggm = _to_float(s_data.get("pv_ggm_tv")) or 0
                pv_exit = _to_float(s_data.get("pv_exit_tv")) or 0
                net_debt = _to_float(s_data.get("net_debt")) or 0
                eq_val = _to_float(s_data.get("equity_value")) or 0

                # 瀑布图：累计计算
                values = [pv_fcf, pv_ggm, pv_exit, -net_debt, eq_val]
                bottoms = []
                heights = []
                cumulative = 0.0
                for i, v in enumerate(values):
                    if i == 0:
                        bottoms.append(0)
                        heights.append(v)
                        cumulative = v
                    elif i == len(values) - 1:
                        bottoms.append(0)
                        heights.append(v)
                    else:
                        bottoms.append(cumulative)
                        heights.append(v)
                        cumulative += v

                offset = (g_idx - n_groups / 2 + 0.5) * bar_width
                xs = [x + offset for x in range(n_bars)]
                bar_colors = []
                for i, h in enumerate(heights):
                    if i == 0:
                        bar_colors.append(colors_scenario[g_idx])
                    elif i == len(heights) - 1:
                        bar_colors.append(colors_scenario[g_idx])
                    else:
                        bar_colors.append(MS_PALETTE[3] if h >= 0 else MS_PALETTE[5])  # Green/Red

                ax.bar(xs, heights, bottom=bottoms, width=bar_width,
                       color=bar_colors, edgecolor='white', linewidth=0,
                       label=s_name, zorder=3)

                # Connector lines between bars within this scenario
                _add_waterfall_connectors(ax, xs, bottoms, heights)

                # 数值标注
                for i, (x, h, b) in enumerate(zip(xs, heights, bottoms)):
                    display_val = h if i in (0, len(heights) - 1) else h
                    sign = "+" if display_val > 0 and i not in (0, len(heights) - 1) else ""
                    if abs(display_val) > 0:
                        lbl_color = colors_scenario[g_idx]
                        if i not in (0, len(heights) - 1):
                            lbl_color = MS_PALETTE[3] if display_val > 0 else MS_PALETTE[5]
                        ax.text(x, b + h + 2, f"{sign}{display_val:.0f}",
                                ha="center", va="bottom", fontsize=7,
                                color=lbl_color, fontweight='bold', zorder=5)

            # Equity Value summary annotation at bottom
            eq_vals = []
            for _, s_data in scenario_list:
                eq_vals.append(_to_float(s_data.get("equity_value")) or 0)
            eq_summary = " / ".join(
                [f"{names[0]}: {sym}{vals:.0f}M"
                 for names, vals in zip(scenario_list, eq_vals)])
            fig.text(0.5, 0.04, f"Equity Value Summary: {eq_summary}",
                     ha='center', va='bottom', fontsize=7,
                     color='#4A4A4A', fontstyle='italic')

            ax.set_xticks(range(n_bars))
            ax.set_xticklabels(labels, fontsize=8)
            ax.set_ylabel(f"Value ({sym}M)", fontsize=9, color=MS_CHART_TICK_COLOR)
            ax.legend(frameon=False, fontsize=8, labelcolor=MS_CHART_TICK_COLOR,
                      loc='upper left')
            # Apply publication style + source watermark
            _apply_ax_publication_style(ax, title="Valuation Bridge by Scenario")
            _add_source_watermark(fig)
            fig.tight_layout(pad=1.5)

            fd, path = tempfile.mkstemp(prefix="ms_bridge_", suffix=".png")
            os.close(fd)
            try:
                fig.savefig(path, dpi=300, facecolor='#FFFFFF',
                            bbox_inches='tight')
                doc.add_picture(path, width=Inches(6.2))
            except Exception:
                pass
            finally:
                plt.close(fig)
                try:
                    os.remove(path)
                except OSError:
                    pass
        except Exception:
            plt.close("all")

    # 分析文字
    doc.add_paragraph()
    bridge_analysis = data.get("bridge_analysis_text")
    if bridge_analysis:
        _body_paragraph(doc, bridge_analysis, theme_colors)
    else:
        base_eq = base.get("equity_value", "—")
        base_fcf = base.get("pv_fcf", "—")
        default_text = (
            f"基准情景下，股权价值为 {sym}{base_eq:,.0f}M，"
            f"其中 FCF 现值贡献 {sym}{base_fcf:,.0f}M。"
            "终端价值（GGM 与 Exit Multiple 取均值）占总估值的主要部分，"
            "体现了长期增长假设对估值的决定性影响。"
        )
        _body_paragraph(doc, default_text, theme_colors)

    doc.add_page_break()


# =============================================================================
# 21. 主入口
# =============================================================================

def make_report(data: Dict[str, Any], output_path: str,
                theme: str = "classic", language: str = "zh") -> str:
    """按 Morgan Stanley 风格生成 docx 研究报告并保存到 ``output_path``。

    参数
    ----
    data: 报告数据字典，字段均为可选。至少 ``title_cn`` 或 ``title_en``
          用于封面，``executive_summary`` / ``key_takeaways`` 用于摘要，
          ``sections`` 列表用于章节内容，``metrics`` / ``financial_table``
          / ``rating_table`` / ``shovel_stocks`` / ``exhibits`` /
          ``appendix`` 等按需提供。新增可选字段：
          ``thesis_charts`` / ``value_chain`` / ``shovel_stocks_list``
          / ``whats_changed`` / ``industry_view`` / ``analyst_title``
          / ``analyst_email`` / ``industry`` / ``entity``。
    output_path: 输出 .docx 路径。上层应确保父目录存在，此处不负责创建。
    theme: 主题名称，目前支持 "classic"（深蓝+金）。
    language: 预留字段，"zh"/"en"/"bilingual"，目前章节标题双语显示。

    返回
    ----
    output_path: 写入成功后返回传入路径。
    """
    theme_colors = _theme_colors(theme)

    # 重置 Exhibit 全局计数器
    _reset_exhibit_counter()

    doc = Document()

    # 基础样式：默认字体与边距
    try:
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    except Exception:
        pass

    # 临时图表资源池，用于统一清理
    chart_temp_paths: List[str] = []

    # 按照规范顺序逐页渲染（缺失字段自动跳过）
    # 1. 封面页（MS 标准模板）
    _cover_page(doc, data, theme_colors)
    # 2. 目录
    _toc_page(doc, data, theme_colors)
    # 3. Our Thesis in N Charts（如果有 thesis_charts 数据）
    _thesis_charts_section(doc, data, theme_colors, chart_temp_paths)
    # 4. 执行摘要
    _executive_summary(doc, data, theme_colors)

    # 5. 正文章节
    sectors_allocation = data.get("sectors_allocation")
    sections_list = list(data.get("sections") or [])
    for idx, section in enumerate(sections_list):
        _section(doc, section, theme_colors,
                 section_index=idx,
                 sectors_allocation=sectors_allocation)
    # 章节后统一换页
    if sections_list:
        doc.add_page_break()

    _kpi_blocks(doc, data, theme_colors)
    _financial_table(doc, data, theme_colors)

    # 财务图表页：当 data 含 financial_chart 时，插入 bar + line + pie + comps 图
    financial_chart = data.get("financial_chart")
    if isinstance(financial_chart, dict) and financial_chart:
        head = doc.add_paragraph(style='Heading 1')
        _add_run(head, "财务图表 / Financial Charts",
                 size=18, bold=True, color=theme_colors["brand_blue"])
        _gold_short_underline(doc, theme_colors)
        doc.add_paragraph()

        # Exhibit A：收入 & EBITDA 双柱图
        ex_a = doc.add_paragraph()
        _add_run(ex_a, _next_exhibit_label("收入与 EBITDA 对比"),
                 size=13, bold=True, color=theme_colors["brand_blue"])
        a_data = {
            "type": "bar",
            "title": "Revenue & EBITDA ($M)",
            "labels": list(financial_chart.get("labels") or []),
            "revenue": list(financial_chart.get("revenue") or []),
            "ebitda": list(financial_chart.get("ebitda") or []),
        }
        if a_data["labels"] and (a_data["revenue"] or a_data["ebitda"]):
            path_a = _add_chart_image(doc, a_data, theme_colors)
            if path_a:
                chart_temp_paths.append(path_a)
        doc.add_paragraph()

        # Exhibit B：利润率折线图
        ex_b = doc.add_paragraph()
        _add_run(ex_b, _next_exhibit_label("利润率趋势"),
                 size=13, bold=True, color=theme_colors["brand_blue"])
        b_data = {
            "type": "line",
            "title": "Margin Trends (%)",
            "labels": list(financial_chart.get("labels") or []),
            "gross_margin": list(financial_chart.get("gross_margin") or []),
            "ebitda_margin": list(financial_chart.get("ebitda_margin") or []),
            "net_margin": list(financial_chart.get("net_margin") or []),
        }
        if b_data["labels"] and (b_data["gross_margin"]
                                  or b_data["ebitda_margin"]
                                  or b_data["net_margin"]):
            path_b = _add_chart_image(doc, b_data, theme_colors)
            if path_b:
                chart_temp_paths.append(path_b)
        doc.add_paragraph()

        # Exhibit C：SOTP 业务板块饼图
        sotp_pie = financial_chart.get("sotp_pie")
        if isinstance(sotp_pie, dict) and sotp_pie.get("labels") and sotp_pie.get("values"):
            ex_c = doc.add_paragraph()
            _add_run(ex_c, _next_exhibit_label("SOTP — Enterprise Value by Segment"),
                     size=13, bold=True, color=theme_colors["brand_blue"])
            c_data = {
                "type": "pie",
                "title": "SOTP - Enterprise Value by Segment",
                "labels": list(sotp_pie.get("labels") or []),
                "values": list(sotp_pie.get("values") or []),
            }
            path_c = _add_chart_image(doc, c_data, theme_colors)
            if path_c:
                chart_temp_paths.append(path_c)
            doc.add_paragraph()

        # Exhibit D：可比公司倍数对比图
        comps_bar = financial_chart.get("comps_bar")
        if isinstance(comps_bar, dict) and comps_bar.get("labels"):
            ex_d = doc.add_paragraph()
            _add_run(ex_d, _next_exhibit_label("可比公司 EV/Revenue 倍数对比"),
                     size=13, bold=True, color=theme_colors["brand_blue"])
            d_data = {
                "type": "bar",
                "title": "Comparable Companies — EV/Revenue Multiples",
                "labels": list(comps_bar.get("labels") or []),
                "revenue": list(comps_bar.get("ev_rev_ltm") or []),
                "ebitda": list(comps_bar.get("ev_rev_ntm") or []),
            }
            path_d = _add_chart_image(doc, d_data, theme_colors)
            if path_d:
                chart_temp_paths.append(path_d)

        doc.add_page_break()

    _rating_table(doc, data, theme_colors)
    _shovel_stocks(doc, data, theme_colors)

    # 6. 产业链分析（如果有 value_chain 数据）
    _value_chain_section(doc, data, theme_colors)
    # 7. 铲子股列表（如果有 shovel_stocks_list 数据）
    _shovel_stocks_list(doc, data, theme_colors)

    # 8. DCF 章节（按顺序：情景分析 → WACC → 敏感性 → 可比公司 → 估值桥）
    _scenario_comparison_section(doc, data, theme_colors, language)
    _wacc_section(doc, data, theme_colors, language)
    _sensitivity_section(doc, data, theme_colors, language)
    _comps_section(doc, data, theme_colors, language)
    _valuation_bridge_section(doc, data, theme_colors, language)

    # 9. Exhibits
    _exhibits_section(doc, data, theme_colors)
    # 10. 附录与披露
    _appendix(doc, data, theme_colors)
    _disclosure(doc, data, theme_colors)

    # 保存：不负责创建父目录
    doc.save(output_path)

    # 清理临时图表 PNG
    _cleanup_chart_paths(chart_temp_paths)

    return output_path


# =============================================================================
# 17. 命令行入口（保持向后兼容的测试入口）
# =============================================================================

if __name__ == "__main__":
    import json
    import sys

    sample = {
        "company_name": "极光资本控股",
        "title_cn": "为下一个10倍科技周期，资本是瓶颈吗？",
        "title_en": "Financing the Next 10x Tech Cycle",
        "subtitle": "关键论点一句话：算力军备竞赛将重塑上游资本开支结构。",
        "rating": "Overweight",
        "target_price": 125.0,
        "current_price": 98.5,
        "date_str": "2026-06-12",
        "analyst": "陈嘉怡 CFA",
        "research_type": "Foundation",
        "currency": "USD",
        "key_takeaways": [
            "算力需求在未来三年预计以 35% CAGR 扩张，资本开支将成为瓶颈。",
            "上游芯片与光模块厂商最受益，利润池份额有望翻倍。",
            "估值溢价将集中于具备专有硅片与供应渠道的企业。",
            "地缘与出口管制风险可能压缩远期回报曲线。",
        ],
        "executive_summary": {
            "thesis": [
                "算力军备竞赛重塑上游资本开支结构，半导体链条议价能力持续上升。",
                "预计 2026-2028 年全球数据中心资本开支总额突破 5000 亿美元。",
                "芯片与光模块份额从 38% 提升至 52%，利润池有望翻倍。",
            ],
            "tp_and_upside": [
                "目标价 $125.0（当前 $98.5），上行空间约 +27%。",
                "评级：Overweight（超配），估值锚定 26E P/E 22×。",
            ],
            "risks": [
                "• 出口管制与地缘冲突加剧导致订单延迟与估值压缩。",
                "• 客户集中度偏高，单一大型云厂商预算调整将显著影响收入。",
                "• 产能瓶颈可能导致毛利率在 2H26 出现阶段性回落。",
            ],
            "catalysts": [
                "• 下一代 AI 加速器量产时间表公布。",
                "• 大型云厂商资本开支指引上调。",
                "• 自研硅片客户的订单结构优化。",
            ],
        },
        "financial_chart": {
            "labels": ["FY24", "FY25E", "FY26E", "FY27E"],
            "revenue": [4810, 5230, 5750, 6520],
            "ebitda": [1210, 1405, 1580, 1820],
            "gross_margin": [52.1, 54.0, 55.5, 56.8],
            "ebitda_margin": [25.2, 26.9, 27.5, 27.9],
            "net_margin": [18.5, 19.8, 21.0, 22.1],
        },
        "sectors_allocation": {
            "title": "行业配置 / Sector Allocation",
            "labels": ["半导体", "光模块", "电源/散热", "IDC", "软件应用", "其他"],
            "values": [40.0, 25.0, 12.0, 10.0, 8.0, 5.0],
        },
        "sections": [
            {
                "title_cn": "一、算力需求：结构性扩张",
                "title_en": "I. Compute Demand — Structural Expansion",
                "paragraphs": [
                    "全球云厂商与大型科技公司的训练集群在 2024-2025 年经历了一轮快速扩充，"
                    "推理侧需求接棒成为新的增长引擎。",
                    "我们估计单位推理成本将在未来两年以每年 30% 的速度下降，"
                    "推动应用层渗透率显著提升。",
                    "• 训练侧：高毛利但订单集中，议价权集中于头部三大云厂商。",
                    "• 推理侧：客户分散，单位推理成本下降打开商业化空间。",
                    "  - 模型即服务（MaaS）平台化趋势。",
                    "  - 端侧推理芯片出货量高速增长。",
                ],
                "subsections": [
                    {
                        "title": "1.1 训练 vs 推理",
                        "paragraphs": [
                            "训练侧高毛利率但订单集中于少数客户，推理侧更分散且对能效比敏感。",
                            "• 训练：每 GPU 年 TCO 仍在 $35k-55k 区间，训练主导 2024-25 采购。",
                            "• 推理：端到端时延要求驱动边缘与专用 ASIC 需求。",
                        ],
                    },
                ],
            },
            {
                "title_cn": "二、资本开支结构重塑",
                "title_en": "II. Capex Structure Reshaping",
                "paragraphs": [
                    "我们预期 2026 年起，数据中心电源、光互联与散热配套将成为新的投资热点。",
                    "• 单机柜功耗从 15 kW 跃升至 30-60 kW。",
                    "• 液冷方案渗透率加速提升。",
                ],
            },
            {
                "title_cn": "三、行业配置与推荐",
                "title_en": "III. Sector Allocation & Recommendations",
                "paragraphs": [
                    "我们建议超配上中游半导体与光模块，标配 IDC 与软件应用，低配传统硬件制造。",
                    "• 风险收益比最高的细分：光模块（800G/1.6T 迭代）。",
                    "• 阿尔法来源：具备自研硅片与稳定产能的厂商。",
                ],
            },
        ],
        "metrics": [
            {"value": "$5.2T", "unit": "TAM 2026E", "label": "市场规模"},
            {"value": "35%", "unit": "CAGR '24-'28", "label": "复合增长"},
            {"value": "+27%", "unit": "Upside", "label": "目标价上涨空间"},
        ],
        "financial_table": {
            "headers": ["指标", "FY24", "FY25E", "FY26E", "FY27E"],
            "rows": [
                ["营业收入 ($M)", "4,810", "5,230", "5,750", "6,520"],
                ["毛利率 (%)", "52.1", "54.0", "55.5", "56.8"],
                ["EBITDA ($M)", "1,210", "1,405", "1,580", "1,820"],
                ["EBITDA 利润率 (%)", "25.2", "26.9", "27.5", "27.9"],
                ["CAPEX ($M)", "-620", "-780", "-910", "-1,050"],
                ["调整后 EPS ($)", "2.12", "2.48", "2.91", "3.38"],
            ],
            "highlight_rows": [5],
            "auto_yoy": True,
            "source": "Source: Morgan Stanley Research. 单位：百万美元 / %。",
        },
        "rating_table": [
            ["Ticker", "Rating", "Last Close", "Company", "Reason"],
            ["NVDA.O", "OW", "198.35", "Nvidia", "AI compute leader"],
            ["AVGO.O", "OW", "82.10", "Broadcom", "Custom ASIC leverage"],
            ["AMD.O", "EW", "162.40", "AMD", "Balanced mix"],
            ["INTC.O", "UW", "31.15", "Intel", "Foundry drag"],
        ],
        "shovel_stocks": [
            {"rank": 1, "company": "Nvidia", "ticker": "NVDA",
             "product": "H200 / B200 GPU", "analyst": "E. Lee",
             "market_cap_mn": 2400000.0, "perf_1y_pct": 112.3},
            {"rank": 2, "company": "Broadcom", "ticker": "AVGO",
             "product": "Custom AI ASIC", "analyst": "K. Wu",
             "market_cap_mn": 680000.0, "perf_1y_pct": 72.5},
            {"rank": 3, "company": "Marvell", "ticker": "MRVL",
             "product": "Data Center PHY", "analyst": "J. Zhang",
             "market_cap_mn": 85000.0, "perf_1y_pct": 23.1},
            {"rank": 4, "company": "Applied Materials", "ticker": "AMAT",
             "product": "先进制程设备", "analyst": "M. Chen",
             "market_cap_mn": 145000.0, "perf_1y_pct": -5.2},
        ],
    }

    out = sys.argv[1] if len(sys.argv) > 1 else "/tmp/ms_test.docx"
    if len(sys.argv) > 2:
        with open(sys.argv[2], "r", encoding="utf-8") as f:
            sample = json.load(f)
    make_report(sample, out)

    # 自测：统计 paragraphs / tables / 粗略页数字段
    doc_for_check = Document(out)
    paragraphs = len(doc_for_check.paragraphs)
    tables = len(doc_for_check.tables)
    # 粗略页数：docx 原生没有 page count；使用 section * 10 + breaks + tables 估算
    page_breaks = sum(1 for p in doc_for_check.paragraphs
                      if p.runs and any(run.text for run in p.runs)
                      and "PAGE BREAK" in str(p._p.xml))
    # 更稳妥：根据文档段落总数粗略估算，按每页 ~20 段落折算
    pages = max(1, (paragraphs + 19) // 20 + page_breaks)
    print(f"DOCX_TEST_OK pages={pages} paragraphs={paragraphs} tables={tables}")
